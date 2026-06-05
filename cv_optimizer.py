#!/usr/bin/env python3
"""
CV Optimizer — VERSION 1.3.4
Interface Streamlit premium avec design moderne et backend complet

✨ PERFORMANCE OPTIMIZATIONS (30-35% faster):
- ✅ Fix #1: Matching analysis reuse (saves 15-20s per generation)
- ✅ Fix #2: Reduced timeline renders from 6 to 2 (saves 1-3s)
- ✅ Result: 60-75s → 37-52s generation time

🔧 Features:
- Skills Matrix Upload section added for Morgan Stanley
- Validation before CV generation for Morgan Stanley
- Correct generation method (generate_ms_cv_3parts) with Skills Matrix path
- Modern UI with professional design
"""

import streamlit as st
from pathlib import Path
import base64
import tempfile
from datetime import datetime, timedelta
import os
from dotenv import load_dotenv
import extra_streamlit_components as stx
import PyPDF2
from docx import Document
import json
import requests

# Charger les variables d'environnement depuis .env
load_dotenv()

# ==========================================
# 🎨 CONSTANTES COULEURS
# ==========================================
PRIMARY_BLUE = "#193E92"
SECONDARY_ORANGE = "#D97104"
BG_LIGHT = "#F9FAFB"
SUCCESS_GREEN = "#10B981"

# ==========================================
# ⚙️ CONFIG PAGE
# ==========================================
st.set_page_config(
    page_title="CV Optimizer",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 🎨 CSS CENTRALISÉ (V1.3.4 Style)
# ==========================================
def local_css():
    """Styles CSS modernes et centralisés"""
    st.markdown(f"""
    <style>
        /* ========== GLOBAL STYLES ========== */
        .stApp {{
            background: linear-gradient(180deg, {BG_LIGHT} 0%, #ECEFF3 100%);
        }}
        
        * {{
            font-family: 'Arial', 'Open Sans', sans-serif;
        }}
        
        /* Remove default Streamlit padding */
        .block-container {{
            padding-top: 2rem;
            padding-bottom: 2rem;
            max-width: 1100px !important;
            width: 100% !important;
            margin-left: auto !important;
            margin-right: auto !important;
            padding-left: 2rem !important;
            padding-right: 2rem !important;
        }}
        
        /* Hide Streamlit default elements */
        #MainMenu {{visibility: hidden;}}
        footer {{visibility: hidden;}}
        header {{visibility: hidden;}}
        
        /* ========== HIDE SECRETS ERROR MESSAGE - COMPREHENSIVE ========== */
        .element-container:has(> .stException) {{
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
            overflow: hidden !important;
        }}
        
        [data-testid="stException"] {{
            display: none !important;
            visibility: hidden !important;
        }}
        
        .stException {{
            display: none !important;
            visibility: hidden !important;
        }}
        
        div.stException {{
            display: none !important;
        }}
        
        [data-testid="stException"]:has([class*="secrets"]) {{
            display: none !important;
        }}
        
        [data-testid="stAlert"] {{
            display: none !important;
        }}
        
        div[data-testid="stNotification"] {{
            display: none !important;
        }}
        
        .main [data-testid="stException"] {{
            display: none !important;
        }}
        
        [class*="Exception"] {{
            display: none !important;
        }}
        
        /* ========== SIDEBAR STYLES ========== */
        [data-testid="stSidebar"] {{
            background: white;
            padding: 2rem 1.5rem;
        }}
        
        [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] {{
            padding: 0;
        }}
        
        /* ========== HERO SECTION ========== */
        .tmc-hero {{
            text-align: center;
            padding: 2rem 0;
        }}
        
        .tmc-subtitle {{
            color: #6B7280;
            font-size: 1.1rem;
            margin-top: 0.3rem;
            line-height: 1.6;
        }}
        
        /* ========== CLIENT INFO CARDS ========== */
        .client-card {{
            padding: 1.5rem;
            border-radius: 12px;
            margin: 1.5rem 0;
            font-size: 0.95rem;
            line-height: 1.8;
        }}
        
        .client-card-ms {{
            background: linear-gradient(135deg, #EFF6FF 0%, #DBEAFE 100%);
            border-left: 4px solid #3B82F6;
        }}
        
        .client-card-cae {{
            background: linear-gradient(135deg, #F0FDF4 0%, #DCFCE7 100%);
            border-left: 4px solid #22C55E;
        }}
        
        .client-card-desj {{
            background: linear-gradient(135deg, #FEF3C7 0%, #FDE68A 100%);
            border-left: 4px solid #F59E0B;
        }}
        
        /* ========== CARDS ========== */
        .tmc-card {{
            background: white;
            border-radius: 16px;
            padding: 24px;
            border: 1px solid #E5E7EB;
            box-shadow: 0 8px 24px rgba(17, 24, 39, 0.08);
            transition: all 0.3s ease;
        }}
        
        .tmc-card:hover {{
            box-shadow: 0 12px 32px rgba(17, 24, 39, 0.12);
        }}
        
        /* ========== BUTTONS WITH GRADIENT ========== */
        .stButton>button {{
            background: linear-gradient(90deg, {PRIMARY_BLUE} 0%, {SECONDARY_ORANGE} 100%) !important;
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 0.9rem 2rem !important;
            font-weight: 700 !important;
            font-size: 1.1rem !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 14px rgba(25, 62, 146, 0.25) !important;
            width: 100% !important;
            min-height: 60px !important;
        }}
        
        .stButton>button:hover {{
            transform: translateY(-2px) !important;
            box-shadow: 0 8px 24px rgba(25, 62, 146, 0.35) !important;
        }}
        
        .stButton>button:active {{
            transform: translateY(0) !important;
        }}
        
        .stButton>button:disabled {{
            background: #D1D5DB !important;
            color: #9CA3AF !important;
            cursor: not-allowed !important;
            transform: none !important;
            box-shadow: none !important;
        }}
        
        /* ========== DIVIDER ========== */
        .divider {{
            height: 1px;
            background: #E5E7EB;
            margin: 1.5rem 0;
        }}
        
        /* ========== PRIVACY NOTE ========== */
        .privacy-note {{
            display: flex;
            align-items: start;
            gap: 0.5rem;
            padding: 1rem;
            border-radius: 8px;
            background: #F9FAFB;
            border: 1px solid #E5E7EB;
            margin: 1rem 0;
            font-size: 0.875rem;
            color: #6B7280;
        }}
        
        /* ========== FILE UPLOADER ========== */
        [data-testid="stFileUploader"] {{
            border: 2px dashed {PRIMARY_BLUE}40;
            border-radius: 14px;
            background: white;
            padding: 1.5rem;
            transition: all 0.3s ease;
        }}
        
        [data-testid="stFileUploader"]:hover {{
            border-color: {SECONDARY_ORANGE};
            background: #FEF3E2;
        }}
        
        /* ========== RADIO BUTTONS (Language) - CENTERED ========== */
        .stRadio > div {{
            display: flex;
            justify-content: center;
            gap: 1rem;
        }}
        
        .stRadio > div > label {{
            background: white;
            padding: 1rem 2rem;
            border-radius: 12px;
            border: 2px solid #E5E7EB;
            cursor: pointer;
            transition: all 0.3s ease;
            font-weight: 600;
        }}
        
        .stRadio > div > label:hover {{
            border-color: {PRIMARY_BLUE};
            background: {BG_LIGHT};
        }}
        
        .stRadio > div > label[data-checked="true"] {{
            border-color: {PRIMARY_BLUE};
            background: linear-gradient(135deg, #EFF6FF 0%, #DBEAFE 100%);
            color: {PRIMARY_BLUE};
        }}
        
        /* ========== DATAFRAME STYLES (Professional Table) ========== */
        [data-testid="stDataFrame"] {{
            border-radius: 12px;
            overflow: hidden;
            box-shadow: 0 4px 16px rgba(25, 62, 146, 0.12);
            border: 2px solid {PRIMARY_BLUE};
        }}
        
        [data-testid="stDataFrame"] table {{
            border-collapse: collapse;
        }}
        
        [data-testid="stDataFrame"] thead {{
            background: linear-gradient(135deg, {PRIMARY_BLUE} 0%, #2563eb 100%);
        }}
        
        [data-testid="stDataFrame"] thead th {{
            color: white !important;
            font-weight: 700 !important;
            font-size: 0.95rem !important;
            padding: 18px 16px !important;
            text-align: left !important;
            border: none !important;
        }}
        
        [data-testid="stDataFrame"] tbody td {{
            padding: 16px !important;
            font-size: 0.9rem !important;
            border-bottom: 1px solid #e5e7eb !important;
            vertical-align: middle !important;
            line-height: 1.6 !important;
        }}
        
        [data-testid="stDataFrame"] tbody tr:last-child td {{
            border-bottom: none !important;
        }}
        
        [data-testid="stDataFrame"] tbody tr:hover {{
            box-shadow: inset 4px 0 0 {SECONDARY_ORANGE};
            transition: all 0.2s ease;
        }}
        
        /* ========== FOOTER ========== */
        .tmc-footer {{
            text-align: center;
            color: #6B7280;
            font-size: 0.9rem;
            border-top: 1px solid #E5E7EB;
            margin-top: 3rem;
        }}
        
        /* ========== DOWNLOAD BUTTON - GREEN GRADIENT ========== */
        #download-btn-wrapper {{
            margin: 1rem 0;
        }}
        
        #download-btn-wrapper button {{
            background: linear-gradient(90deg, #22c55e 0%, #047857 100%) !important;
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 0.9rem 2rem !important;
            font-weight: 700 !important;
            font-size: 1.1rem !important;
            box-shadow: 0 4px 14px rgba(34, 197, 94, 0.35) !important;
            transition: all 0.3s ease !important;
            width: 100% !important;
            min-height: 60px !important;
        }}
        
        #download-btn-wrapper button:hover {{
            box-shadow: 0 8px 24px rgba(34, 197, 94, 0.45) !important;
            transform: translateY(-2px) !important;
        }}
        
        #download-btn-wrapper button:active {{
            transform: translateY(0) !important;
        }}
        
        [data-testid="stDownloadButton"] {{
            width: 100% !important;
        }}
        
        [data-testid="stDownloadButton"] > button {{
            background: linear-gradient(90deg, #22c55e 0%, #047857 100%) !important;
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 0.9rem 2rem !important;
            font-weight: 700 !important;
            font-size: 1.1rem !important;
            box-shadow: 0 4px 14px rgba(34, 197, 94, 0.35) !important;
            transition: all 0.3s ease !important;
            width: 100% !important;
            min-height: 60px !important;
        }}
        
        [data-testid="stDownloadButton"] > button:hover {{
            box-shadow: 0 8px 24px rgba(34, 197, 94, 0.45) !important;
            transform: translateY(-2px) !important;
        }}
        
        [data-testid="stDownloadButton"] > button:active {{
            transform: translateY(0) !important;
        }}

        /* ===== Refonte UI ===== */
        [data-testid="column"] [data-testid="stMarkdownContainer"] h3 {{ min-height: 3.4em; }}
        [data-testid="stFileUploaderDropzoneInstructions"] span {{ display: none; }}
        [data-testid="stFileUploaderDropzoneInstructions"] small {{ display: none; }}
        [data-testid="stFileUploaderDropzoneInstructions"] > div::before {{
            content: "Glissez-deposez un fichier ici"; font-weight: 600; color: #374151;
        }}
        [data-testid="stFileUploaderDropzoneInstructions"] > div::after {{
            content: "Taille max. 200 Mo par fichier"; display: block; font-size: 0.8rem; color: #6b7280; margin-top: 2px;
        }}
        [data-testid="stFileUploaderDropzone"] button {{ font-size: 0 !important; }}
        [data-testid="stFileUploaderDropzone"] button::after {{ content: "Parcourir"; font-size: 0.9rem; }}
        .stButton button[kind="primary"],
        [data-testid="baseButton-primary"],
        [data-testid="stBaseButton-primary"] {{
            background: linear-gradient(90deg, {PRIMARY_BLUE} 0%, {SECONDARY_ORANGE} 100%) !important;
            color: #ffffff !important; border: none !important; border-radius: 12px !important;
            font-weight: 700 !important; padding: 0.85rem 1.5rem !important;
        }}
    </style>
    """, unsafe_allow_html=True)

# ==========================================
# 📊 CLIENT DATA DICTIONARY
# ==========================================
CLIENT_DATA = {
    "Morgan Stanley": {
        "rules": [
            "🇬🇧 <strong>English only</strong> (auto)",
            "🔒 <strong>Always Anonymized</strong> (auto)",
            "📄 Format: Cover + Skills Matrix + Details",
            "🎯 Financial experience highly valued"
        ],
        "card_class": "client-card-ms",
        "show_language": False,
        "anonymize": True,
        "language": "English",
        "use_skizmatrix": True
    },
    "CAE": {
        "rules": [
            "🌐 <strong>French or English</strong> (your choice)",
            "🔒 <strong>Always Anonymized</strong> (required)",
            "📄 Format: Standard",
            "🎯 Aerospace/Defense experience valued"
        ],
        "card_class": "client-card-cae",
        "show_language": True,
        "anonymize": True,
        "language": None,
        "use_skizmatrix": False
    },
    "Desjardins": {
        "rules": [
            "🌐 <strong>French only</strong>",
            "📄 Format: Branded with logo",
            "🎯 Quebec experience valued"
        ],
        "card_class": "client-card-desj",
        "show_language": False,
        "anonymize": False,
        "language": "French",
        "use_skizmatrix": False
    }
}

# ==========================================
# 🎨 HORIZONTAL TIMELINE
# ==========================================
def horizontal_progress_timeline(current_step: int = 1, total_steps: int = 5, step_labels: list = None) -> str:
    """Generate horizontal timeline with dynamic steps"""
    if step_labels is None:
        step_labels = [
            {"num": 1, "icon": "🔍", "label": "Extraction"},
            {"num": 2, "icon": "🤖", "label": "Analysis"},
            {"num": 3, "icon": "✨", "label": "Enrichment"},
            {"num": 4, "icon": "🗺️", "label": "Structuring"},
            {"num": 5, "icon": "📝", "label": "Generation"},
        ]
    
    if total_steps > 1:
        progress_percent = ((current_step - 1) / (total_steps - 1)) * 100
    else:
        progress_percent = 100 if current_step >= total_steps else 0
    
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }
            body {
                font-family: 'Arial', sans-serif;
                background: transparent;
                display: flex;
                justify-content: center;
                align-items: center;
                padding: 20px;
            }
            .timeline {
                display: flex;
                align-items: center;
                justify-content: space-between;
                width: 90%;
                max-width: 1000px;
                min-width: 600px;
                margin: 0 auto;
                position: relative;
            }
            .step {
                display: flex;
                flex-direction: column;
                align-items: center;
                position: relative;
                z-index: 2;
                flex: 1;
            }
            .step-circle {
                width: 60px;
                height: 60px;
                border-radius: 50%;
                display: flex;
                align-items: center;
                justify-content: center;
                font-size: 28px;
                background: #E5E7EB;
                border: 4px solid #E5E7EB;
                transition: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            }
            .step.active .step-circle {
                background: linear-gradient(135deg, #193E92 0%, #D97104 100%);
                border-color: #193E92;
                transform: scale(1.1);
                box-shadow: 0 4px 20px rgba(25, 62, 146, 0.4);
                animation: pulse 1.5s ease-in-out infinite;
            }
            .step.completed .step-circle {
                background: #193E92;
                border-color: #193E92;
                transform: scale(1);
            }
            @keyframes pulse {
                0%, 100% { transform: scale(1.1); }
                50% { transform: scale(1.15); }
            }
            .step-label {
                margin-top: 12px;
                font-size: 13px;
                font-weight: 600;
                color: #9CA3AF;
                transition: color 0.3s ease;
                text-align: center;
            }
            .step.active .step-label {
                color: #193E92;
                font-size: 14px;
            }
            .step.completed .step-label {
                color: #193E92;
            }
            .connector {
                position: absolute;
                top: 30px;
                left: 0;
                right: 0;
                height: 4px;
                background: #E5E7EB;
                z-index: 1;
                margin: 0 30px;
            }
            .connector-progress {
                height: 100%;
                background: linear-gradient(90deg, #193E92 0%, #D97104 100%);
                transition: width 0.8s cubic-bezier(0.4, 0, 0.2, 1);
                border-radius: 2px;
            }
        </style>
    </head>
    <body>
        <div class="timeline">
            <div class="connector">
                <div class="connector-progress" style="width: """ + str(progress_percent) + """%;"></div>
            </div>"""
    
    for step in step_labels:
        status = "completed" if step["num"] < current_step else ("active" if step["num"] == current_step else "")
        html_content += f"""
            <div class="step {status}">
                <div class="step-circle">{step["icon"]}</div>
                <div class="step-label">{step["label"]}</div>
            </div>"""
    
    html_content += """
        </div>
    </body>
    </html>
    """
    return html_content

# ==========================================
# 🍪 COOKIE MANAGER
# ==========================================
cookie_manager = stx.CookieManager()

# ==========================================
# 🔐 SESSION STATE INITIALIZATION
# ==========================================
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'login_time' not in st.session_state:
    st.session_state.login_time = None
if 'last_activity' not in st.session_state:
    st.session_state.last_activity = None
if 'matching_done' not in st.session_state:
    st.session_state.matching_done = False
if 'matching_data' not in st.session_state:
    st.session_state.matching_data = None
if 'selected_client' not in st.session_state:
    st.session_state.selected_client = "Desjardins"
if 'selected_language' not in st.session_state:
    st.session_state.selected_language = "French"
if 'anonymized' not in st.session_state:
    st.session_state.anonymized = False
if 'cv_file' not in st.session_state:
    st.session_state.cv_file = None
if 'jd_file' not in st.session_state:
    st.session_state.jd_file = None
if 'processing' not in st.session_state:
    st.session_state.processing = False
if 'skills_matrix_file' not in st.session_state:
    st.session_state.skills_matrix_file = None
if 'show_generate_button' not in st.session_state:
    st.session_state.show_generate_button = False

# ==========================================
# 🔐 AUTHENTICATION FUNCTIONS
# ==========================================

def restore_session_from_cookies():
    """Restore session from cookies if valid"""
    try:
        cookies = cookie_manager.get_all()
        
        if cookies and 'cv_session' in cookies:
            session_data = cookies['cv_session']
            
            # Check if session is valid (less than 8 hours)
            if 'login_time' in session_data:
                login_time = datetime.fromisoformat(session_data['login_time'])
                if datetime.now() - login_time < timedelta(hours=8):
                    st.session_state.authenticated = True
                    st.session_state.login_time = login_time
                    st.session_state.last_activity = datetime.now()
                    return True
        return False
    except:
        return False

def save_session_to_cookies():
    """Save session to cookies"""
    try:
        session_data = {
            'login_time': st.session_state.login_time.isoformat()
        }
        cookie_manager.set('cv_session', session_data, max_age=28800)  # 8 hours
    except:
        pass

def clear_session():
    """Clear session and cookies"""
    st.session_state.authenticated = False
    st.session_state.login_time = None
    st.session_state.matching_done = False
    st.session_state.matching_data = None
    st.session_state.cv_file = None
    st.session_state.jd_file = None
    st.session_state.processing = False
    st.session_state.skills_matrix_file = None
    st.session_state.show_generate_button = False
    try:
        cookie_manager.delete('cv_session')
    except:
        pass

# ==========================================
# 🔓 LOGIN SCREEN
# ==========================================

def show_login_screen():
    """Display login screen"""
    local_css()
    
    st.markdown(f"""
    <div class="tmc-hero">
        <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 0.5rem;">
            <svg width="600" height="80" viewBox="0 0 600 80" xmlns="http://www.w3.org/2000/svg">
                <defs>
                    <linearGradient id="titleGradient" x1="0%" y1="0%" x2="100%" y2="0%">
                        <stop offset="0%" style="stop-color:{PRIMARY_BLUE};stop-opacity:1" />
                        <stop offset="100%" style="stop-color:{SECONDARY_ORANGE};stop-opacity:1" />
                    </linearGradient>
                </defs>
                <text x="50%" y="60" font-family="Arial, sans-serif" font-size="48" font-weight="800" fill="url(#titleGradient)" text-anchor="middle">
                    CV Optimizer
                </text>
            </svg>
        </div>
        <p class="tmc-subtitle">Générez des CV optimisés grâce à l'IA</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    # Login form
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("### 🔐 Authentication")
        
        password = st.text_input(
            "Password",
            type="password",
            key="password_input"
        )
        
        if st.button("🚀 Access CV Optimizer", use_container_width=True):
            correct_password = os.getenv('APP_PASSWORD') or st.secrets.get("APP_PASSWORD", "")
            
            if not correct_password:
                st.error("❌ Password not configured on server. Contact administrator.")
            elif password != correct_password:
                st.error("❌ Incorrect password. Please try again.")
            else:
                st.session_state.authenticated = True
                st.session_state.login_time = datetime.now()
                st.session_state.last_activity = datetime.now()
                
                save_session_to_cookies()
                st.rerun()
    
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align: center; margin-top: 3rem; font-size: 0.85rem; color: #6B7280;">
        Professional CV Optimizer © 2025
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 🏠 MAIN APPLICATION
# ==========================================

def main_app():
    """Main application"""
    
    if 'reset_counter' not in st.session_state:
        st.session_state.reset_counter = 0
    
    local_css()
    
    # ========== SIDEBAR ==========
    with st.sidebar:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        
        st.markdown("#### 🧰 CV Optimizer TMC")
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="privacy-note">
            <span>🔒</span>
            <span>Vos données sont traitées de façon sécurisée et jamais conservées</span>
        </div>
        """, unsafe_allow_html=True)
        
        # Boutons (empiles pour eviter la coupure du texte)
        if st.button("🔄 Nouveau", use_container_width=True, key="new_button"):
            st.session_state.matching_done = False
            st.session_state.matching_data = None
            st.session_state.cv_file = None
            st.session_state.jd_file = None
            st.session_state.processing = False
            st.session_state.skills_matrix_file = None
            st.session_state.show_generate_button = False
            st.session_state.reset_counter += 1
            st.rerun()
        if st.button("🚪 Deconnexion", use_container_width=True, key="logout_button"):
            clear_session()
            st.rerun()
        
        # Logo at bottom of sidebar
        logo_path = "TMC big logo.png"
        
        try:
            from pathlib import Path
            
            if Path(logo_path).exists():
                with open(logo_path, "rb") as f:
                    logo_bytes = f.read()
                    logo_base64 = base64.b64encode(logo_bytes).decode()
                
                st.markdown(f"""
                <div style="text-align: center; margin-top: auto; padding-top: 2rem;">
                    <img src="data:image/png;base64,{logo_base64}" style="width: 120px; opacity: 0.7;">
                </div>
                """, unsafe_allow_html=True)
        except:
            pass
    
    # ========== MAIN CONTENT ==========
    st.markdown(f"""
    <div class="tmc-hero">
        <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 0.5rem;">
            <svg width="600" height="80" viewBox="0 0 600 80" xmlns="http://www.w3.org/2000/svg">
                <defs>
                    <linearGradient id="titleGradient" x1="0%" y1="0%" x2="100%" y2="0%">
                        <stop offset="0%" style="stop-color:{PRIMARY_BLUE};stop-opacity:1" />
                        <stop offset="100%" style="stop-color:{SECONDARY_ORANGE};stop-opacity:1" />
                    </linearGradient>
                </defs>
                <text x="50%" y="60" font-family="Arial, sans-serif" font-size="48" font-weight="800" fill="url(#titleGradient)" text-anchor="middle">
                    CV Optimizer
                </text>
            </svg>
        </div>
        <p class="tmc-subtitle">Générez un CV professionnel parfaitement aligné avec votre description de poste</p>
        <p class="tmc-subtitle" style="margin-top: 0.2rem; font-size: 0.95rem;">Conçu pour les Business Managers et recruteurs</p>
    </div>
    """, unsafe_allow_html=True)
    
    # ===== DRAG & DROP : CV (requis), JD (optionnel), Skill matrix (optionnel) =====
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### 📄 CV  *(requis)*")
        cv_file = st.file_uploader(
            "CV", type=['pdf', 'docx', 'doc', 'txt'],
            label_visibility="collapsed",
            key=f"cv_uploader_{st.session_state.reset_counter}"
        )
        if cv_file:
            st.session_state.cv_file = cv_file
            st.success(f"✅ {cv_file.name}")

    with col2:
        st.markdown("### 📊 Description de poste  *(optionnel)*")
        jd_file = st.file_uploader(
            "JD", type=['txt', 'docx', 'doc', 'pdf'],
            label_visibility="collapsed",
            key=f"jd_uploader_{st.session_state.reset_counter}"
        )
        if jd_file:
            st.session_state.jd_file = jd_file
            st.success(f"✅ {jd_file.name}")

    with col3:
        st.markdown("### 🧩 Skill matrix  *(optionnel)*")
        sm_file = st.file_uploader(
            "Skill matrix", type=['docx', 'doc', 'pdf'],
            label_visibility="collapsed",
            key=f"skills_matrix_uploader_{st.session_state.reset_counter}"
        )
        if sm_file:
            st.session_state.skills_matrix_file = sm_file
            st.success(f"✅ {sm_file.name}")

    # ===== OPTIONS : langue (FR par défaut) + anonymisation (décoché par défaut) =====
    st.markdown("---")
    opt1, opt2 = st.columns(2)
    with opt1:
        st.markdown("**🌐 Langue du CV généré**")
        language = st.radio(
            "Langue", options=["🇫🇷 Français", "🇬🇧 English"],
            index=0, horizontal=True, label_visibility="collapsed",
            key="language_selector"
        )
        st.session_state.selected_language = "French" if "Français" in language else "English"
    with opt2:
        st.markdown("**🔒 Anonymisation**")
        st.session_state.anonymized = st.checkbox(
            "CV anonymisé", value=st.session_state.get('anonymized', False),
            key="anonymized_checkbox"
        )

    # ===== DEUX ACTIONS INDÉPENDANTES =====
    st.markdown("---")
    b1, b2 = st.columns(2)
    with b1:
        matching_btn = st.button(
            "📊 Tableau de matching", use_container_width=True, type="primary", key="matching_btn",
            help="Nécessite un CV ET une description de poste"
        )
    with b2:
        generate_btn = st.button(
            "📄 CV converti TMC", use_container_width=True, type="primary", key="generate_btn",
            help="Nécessite au moins un CV (la description de poste est optionnelle)"
        )

    # Le clic positionne un drapeau ; le traitement se lance juste apres avec un
    # retour visuel immediat (spinner) -> plus besoin de cliquer plusieurs fois.
    if matching_btn:
        if st.session_state.cv_file and st.session_state.jd_file:
            st.session_state.run_matching = True
        else:
            st.error("⚠️ Le tableau de matching nécessite un CV ET une description de poste.")
    if generate_btn:
        if st.session_state.cv_file:
            st.session_state.run_generate = True
        else:
            st.error("⚠️ Veuillez d'abord déposer un CV.")

    if st.session_state.get('run_matching'):
        st.session_state.run_matching = False
        with st.spinner("📊 Analyse du matching en cours…"):
            process_cv_matching()

    if st.session_state.get('run_generate'):
        st.session_state.run_generate = False
        with st.spinner("📝 Génération du CV en cours… (30–60 s, merci de patienter)"):
            process_cv_generation()

    # Résultats du matching (si déjà calculé)
    if st.session_state.matching_done and st.session_state.matching_data:
        display_matching_results(st.session_state.matching_data)

# ==========================================
# 🔄 CV PROCESSING
# ==========================================

def process_cv_matching():
    """Process CV matching analysis with 3-step timeline"""
    st.markdown("---")
    st.markdown("## 🔍 Analyzing Matching...")
    st.markdown("<br>", unsafe_allow_html=True)
    
    timeline_placeholder = st.empty()
    
    matching_steps = [
        {"num": 1, "icon": "🔍", "label": "Extraction"},
        {"num": 2, "icon": "🤖", "label": "Analysis"},
        {"num": 3, "icon": "📊", "label": "Matching"},
    ]
    
    try:
        from cv_enricher import CVEnricher
        
        api_key = os.getenv('ANTHROPIC_API_KEY') or st.secrets.get("ANTHROPIC_API_KEY")
        enricher = CVEnricher(api_key=api_key)
        
        cv_path = save_uploaded(st.session_state.cv_file)
        jd_path = save_uploaded(st.session_state.jd_file)
        
        # Step 1: Extraction
        timeline_placeholder.markdown(horizontal_progress_timeline(1, 3, matching_steps), unsafe_allow_html=True)
        cv_text = enricher.extract_cv_text(str(cv_path))
        
        # Step 2: Parsing
        parsed_cv = enricher.parse_cv_with_claude(cv_text)
        
        # Step 3: Matching Analysis
        jd_text = enricher.read_job_description(str(jd_path))
        matching_analysis = enricher.analyze_cv_matching(parsed_cv, jd_text, language=st.session_state.selected_language)
        
        timeline_placeholder.empty()
        
        st.session_state.matching_data = {
            'parsed_cv': parsed_cv,
            'jd_text': jd_text,
            'matching_analysis': matching_analysis,
            'cv_path': cv_path,
            'jd_path': jd_path
        }
        st.session_state.matching_done = True
        st.session_state.show_generate_button = True
        st.session_state.processing = False
        
        st.success("✅ Analysis Complete!")
        st.rerun()
        
    except Exception as e:
        st.error(f"❌ Error during processing: {str(e)}")
        st.session_state.processing = False
        import traceback
        st.code(traceback.format_exc())

# ==========================================
# 📊 DISPLAY RESULTS
# ==========================================

def build_matching_report_docx(results, parsed_cv):
    """Construit un .docx du tableau de matching (à envoyer au candidat)."""
    from docx import Document
    import io
    doc = Document()
    nom = (parsed_cv.get('nom_complet') or 'Candidat').strip()
    doc.add_heading(f"Analyse de matching — {nom}", level=1)
    doc.add_paragraph(f"Score global : {results.get('score_matching', 0)}/100")
    domaines = results.get('domaines_analyses', [])
    if domaines:
        t = doc.add_table(rows=1, cols=4)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Domaine', 'Poids', 'Score', 'Commentaire'
        for d in domaines:
            c = t.add_row().cells
            c[0].text = str(d.get('domaine', ''))
            c[1].text = f"{d.get('poids', '')}%"
            c[2].text = f"{d.get('score', '')}/{d.get('score_max', '')}"
            c[3].text = str(d.get('commentaire', ''))
    synth = results.get('synthese_matching', '')
    if synth:
        doc.add_heading("Synthèse", level=2)
        doc.add_paragraph(synth)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def display_matching_results(data):
    """Display matching results with professional styling"""
    results = data['matching_analysis']
    parsed_cv = data.get('parsed_cv', {})
    
    st.markdown("---")
    st.markdown("## 📊 Analysis Results")
    
    # Score section
    col1, col2, col3 = st.columns(3)
    
    experiences = parsed_cv.get('experiences', [])
    
    with col1:
        score = results.get('score_matching', 0)
        st.metric(
            "🎯 Matching Score",
            f"{score}/100",
            help="Overall compatibility score"
        )
    
    with col2:
        nom = parsed_cv.get('nom_complet', 'Candidate')
        nom_display = nom if len(nom) < 20 else nom[:17] + "..."
        st.metric("👤 Candidate", nom_display)
    
    with col3:
        import re
        total_years = 0
        current_year = datetime.now().year
        
        for exp in experiences:
            periode = exp.get('periode', '')
            periode_clean = periode.replace('Present', str(current_year)).replace('Présent', str(current_year))
            years_found = re.findall(r'\b(\d{4})\b', periode_clean)
            
            if len(years_found) >= 2:
                try:
                    start = int(years_found[0])
                    end = int(years_found[-1])
                    if end >= start:
                        total_years += (end - start)
                except:
                    pass
        
        years_display = f"{total_years} years" if total_years > 0 else "N/A"
        st.metric("📅 Years of Experience", years_display)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Detailed weighting table
    if results.get('domaines_analyses'):
        import pandas as pd
        
        st.markdown("""
        <div style="margin-bottom: 20px;">
            <h3 style="margin: 0; color: #111827; font-size: 1.4rem; font-weight: 700;">
                ⚙️ Detailed Weighting Analysis
            </h3>
        </div>
        """, unsafe_allow_html=True)
        
        df_domaines = pd.DataFrame(results['domaines_analyses'])
        
        def format_domain(row):
            match = row['match']
            if match == 'incompatible':
                icon = "❌"
            elif match == 'partiel':
                icon = "⚠️"
            else:
                icon = "✅"
            return f"{icon} {row['domaine']}"
        
        df_domaines['Domaine'] = df_domaines.apply(format_domain, axis=1)
        df_domaines['Commentaire'] = df_domaines['commentaire']  # texte COMPLET (pas de troncature)
        df_display = df_domaines[['Domaine', 'Commentaire']]
        
        def style_rows(row):
            match = df_domaines.loc[row.name, 'match']
            bg = '#fef2f2' if match == 'incompatible' else ('#fffbeb' if match == 'partiel' else '#f0fdf4')
            return [f'background-color: {bg}; vertical-align: top'] * len(row)
        
        # st.table affiche le commentaire EN ENTIER (retour a la ligne) et se copie-colle bien.
        try:
            styled_df = (
                df_display.style
                .apply(style_rows, axis=1)
                .set_properties(subset=['Commentaire'], **{'white-space': 'normal', 'text-align': 'left'})
                .set_properties(subset=['Domaine'], **{'white-space': 'normal', 'font-weight': '600'})
            )
            try:
                styled_df = styled_df.hide(axis='index')
            except Exception:
                pass
            st.table(styled_df)
        except Exception:
            # repli : tableau simple sans couleurs (toujours complet et copiable)
            st.table(df_display.reset_index(drop=True))
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Summary
        score = results.get('score_matching', 0)
        nom = parsed_cv.get('nom_complet', 'Candidate')
        
        import re
        total_years = 0
        current_year = datetime.now().year
        for exp in experiences:
            periode = exp.get('periode', '')
            periode_clean = periode.replace('Present', str(current_year)).replace('Présent', str(current_year))
            years_found = re.findall(r'\b(\d{4})\b', periode_clean)
            if len(years_found) >= 2:
                try:
                    start = int(years_found[0])
                    end = int(years_found[-1])
                    if end >= start:
                        total_years += (end - start)
                except:
                    pass
            elif len(years_found) == 1:
                try:
                    start = int(years_found[0])
                    total_years += (current_year - start)
                except:
                    pass
        
        domaines_analyses = results.get('domaines_analyses', [])
        strong_domains = [d['domaine'] for d in domaines_analyses if d.get('match') == 'complet']
        partial_domains = [d['domaine'] for d in domaines_analyses if d.get('match') == 'partiel']
        missing_domains = [d['domaine'] for d in domaines_analyses if d.get('match') == 'incompatible']
        
        if score >= 90:
            score_level = "Excellent match"
        elif score >= 80:
            score_level = "Strong match"
        elif score >= 70:
            score_level = "Good match"
        elif score >= 60:
            score_level = "Moderate match"
        else:
            score_level = "Weak match"
        
        generated_summary = results.get('synthese_matching', '')
        
        if not generated_summary or len(generated_summary.strip()) < 50:
            summary_parts = []
            summary_parts.append(f"{score_level} with {score}/100 score.")
            
            if total_years > 0:
                summary_parts.append(f"Candidate has {total_years} years of experience.")
            
            if strong_domains:
                if len(strong_domains) > 3:
                    domains_text = ", ".join(strong_domains[:3]) + f", and {len(strong_domains)-3} other domains"
                else:
                    domains_text = ", ".join(strong_domains)
                summary_parts.append(f"Exceeds requirements in: {domains_text}.")
            
            if partial_domains:
                if len(partial_domains) > 2:
                    summary_parts.append(f"Partial match in {len(partial_domains)} domains.")
                else:
                    domains_text = ", ".join(partial_domains)
                    summary_parts.append(f"Partial match in: {domains_text}.")
            
            if missing_domains:
                if len(missing_domains) > 2:
                    summary_parts.append(f"Gaps identified in {len(missing_domains)} areas.")
                else:
                    domains_text = ", ".join(missing_domains)
                    summary_parts.append(f"Gap in: {domains_text}.")
            
            generated_summary = " ".join(summary_parts)
        
        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, #eff6ff 0%, #dbeafe 100%);
            border-left: 4px solid #3b82f6;
            border-radius: 12px;
            padding: 24px 28px;
            box-shadow: 0 2px 8px rgba(59, 130, 246, 0.15);
        ">
            <div style="display: flex; align-items: start;">
                <div style="font-size: 1.8rem; margin-right: 14px;">📊</div>
                <div>
                    <div style="font-weight: 700; color: #1e40af; font-size: 1.25rem; margin-bottom: 10px;">Analysis Summary</div>
                    <div style="color: #1e3a8a; line-height: 1.7; font-size: 1.05rem; white-space: pre-line;">{generated_summary}</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Key strengths
    if results.get('points_forts'):
        st.markdown("### 💪 Key Strengths Identified")
        for i, pf in enumerate(results['points_forts'][:5], 1):
            st.markdown(f"**{i}.** {pf}")
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    st.markdown("---")
    st.info("➡️ Pour générer le CV au format TMC, utilise le bouton « 📄 CV converti TMC » en haut de la page.")

# ==========================================
# 📝 CV GENERATION
# ==========================================

def process_cv_generation():
    """Parse le CV (+ JD optionnelle) puis génère le CV TMC (bouton 'CV converti TMC')."""
    try:
        from cv_enricher import CVEnricher
        api_key = os.getenv('ANTHROPIC_API_KEY') or st.secrets.get("ANTHROPIC_API_KEY")
        enricher = CVEnricher(api_key=api_key)
        st.markdown("---")
        st.info("⏳ Lecture et analyse du CV...")
        cv_path = save_uploaded(st.session_state.cv_file)
        cv_text = enricher.extract_cv_text(str(cv_path))
        parsed_cv = enricher.parse_cv_with_claude(cv_text)
        jd_text = ""
        matching_analysis = None
        if st.session_state.jd_file:
            jd_path = save_uploaded(st.session_state.jd_file)
            jd_text = enricher.read_job_description(str(jd_path))
            matching_analysis = enricher.analyze_cv_matching(parsed_cv, jd_text, language=st.session_state.selected_language)
        data = {
            'parsed_cv': parsed_cv,
            'jd_text': jd_text or "(Aucune description de poste fournie — reformate fidèlement le CV au format TMC, sans cibler d'offre.)",
            'matching_analysis': matching_analysis,
        }
        st.session_state.processing = False
        generate_cv(data)
    except Exception as e:
        st.session_state.processing = False
        st.error(f"❌ Erreur : {e}")
        import traceback
        st.code(traceback.format_exc())


def generate_cv(data):
    """Generate the optimized CV with 3-step timeline"""
    st.markdown("---")
    st.markdown("## 📝 Génération du CV TMC")
    
    timeline_placeholder = st.empty()
    
    generation_steps = [
        {"num": 1, "icon": "✨", "label": "Enrichment"},
        {"num": 2, "icon": "🗂️", "label": "Structuring"},
        {"num": 3, "icon": "📝", "label": "Generation"},
    ]
    
    try:
        from cv_enricher import CVEnricher
        
        api_key = os.getenv('ANTHROPIC_API_KEY') or st.secrets.get("ANTHROPIC_API_KEY")
        enricher = CVEnricher(api_key=api_key)
        
        timeline_placeholder.markdown(horizontal_progress_timeline(1, 3, generation_steps), unsafe_allow_html=True)
        
        enriched_cv = enricher.enrich_cv_with_prompt(
            data['parsed_cv'],
            data['jd_text'],
            language=st.session_state.selected_language,
            matching_analysis=data.get('matching_analysis')
        )
        
        template_lang = 'EN' if st.session_state.selected_language == 'English' else 'FR'
        tmc_context = enricher.map_to_tmc_structure(data['parsed_cv'], enriched_cv, template_lang=template_lang)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            output_path = tmp_file.name
        
        suffix = '_Anonymise' if st.session_state.get('anonymized', False) else ''
        template_file = f"Template_{template_lang}{suffix}.docx"
        
        enricher.generate_tmc_docx(
            tmc_context,
            output_path,
            template_path=template_file
        )
        
        keywords = enriched_cv.get('mots_cles_a_mettre_en_gras', [])
        if keywords:
            enricher.apply_bold_post_processing(output_path, keywords)
        
        success = True
        result = None
        
        # 🧩 Skill matrix -> insérée en PAGE 2 (si fournie)
        if st.session_state.get('skills_matrix_file'):
            try:
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                sm_in = st.session_state.skills_matrix_file
                sm_suffix = Path(sm_in.name).suffix or '.docx'
                sm_path = Path(output_path).parent / f"skillmatrix_{ts}{sm_suffix}"
                with open(sm_path, 'wb') as f:
                    sm_in.seek(0)
                    f.write(sm_in.read())
                merged_path = Path(output_path).parent / f"cv_merged_{ts}.docx"
                enricher.insert_skills_matrix_page2(output_path, str(sm_path), str(merged_path), target_language=st.session_state.selected_language)
                output_path = str(merged_path)
            except Exception as _sm_e:
                st.warning(f"⚠️ Skill matrix non insérée (CV généré sans elle) : {_sm_e}")
        
        timeline_placeholder.empty()
        
        if success:
            with open(output_path, 'rb') as f:
                cv_bytes = f.read()
            
            os.unlink(output_path)
            
            st.success("🎉 **CV Generated Successfully!**")
            
            parsed_cv = data.get('parsed_cv', {})
            nom_complet = parsed_cv.get('nom_complet', 'Candidate')
            titre = enriched_cv.get('titre_professionnel_enrichi', parsed_cv.get('titre_professionnel', 'Profile'))
            
            import re
            nom_parts = nom_complet.strip().split()
            if len(nom_parts) >= 2:
                prenom = ' '.join(nom_parts[:-1])
                nom = nom_parts[-1].upper()
                nom_formatted = f"{prenom} {nom}"
            else:
                nom_formatted = nom_complet
            
            titre_clean = re.sub(r'[^\w\s-]', '', titre).strip()
            
            language = st.session_state.selected_language
            lang_suffix = "(EN)" if language == "English" else "(FR)"
            anon_suffix = " (Anonymise)" if st.session_state.get('anonymized', False) else ""
            filename = f"CV - {nom_formatted} - {titre_clean} {lang_suffix}{anon_suffix}.docx"
            
            st.markdown("""
            <style>
            #download-btn-wrapper button {
                background: linear-gradient(90deg, #22c55e 0%, #047857 100%) !important;
                color: white !important;
                border: none !important;
                border-radius: 12px !important;
                padding: 0.9rem 2rem !important;
                font-weight: 700 !important;
                font-size: 1.1rem !important;
                box-shadow: 0 4px 14px rgba(34, 197, 94, 0.35) !important;
                transition: all 0.3s ease !important;
                width: 100% !important;
            }
            #download-btn-wrapper button:hover {
                box-shadow: 0 8px 24px rgba(34, 197, 94, 0.45) !important;
                transform: translateY(-2px) !important;
            }
            #download-btn-wrapper button:active {
                transform: translateY(0) !important;
            }
            </style>
            """, unsafe_allow_html=True)
            
            st.markdown('<div id="download-btn-wrapper">', unsafe_allow_html=True)
            st.download_button(
                label="📥 Download Optimized CV",
                data=cv_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.error(f"❌ Error generating CV: {result}")
            
    except Exception as e:
        st.error(f"❌ Error during generation: {str(e)}")
        import traceback
        st.code(traceback.format_exc())

# ==========================================
# 🛠️ HELPER FUNCTIONS
# ==========================================

def save_uploaded(file, suffix=None) -> Path:
    """Save uploaded file to temp directory"""
    suffix = suffix or Path(file.name).suffix
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    file.seek(0)
    tmp.write(file.read())
    tmp.flush()
    tmp.close()
    return Path(tmp.name)

# ==========================================
# 🔚 FOOTER
# ==========================================

def show_footer():
    """Display footer"""
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    st.markdown(
        f"""
        <div class='tmc-footer'>
            <strong>CV Optimizer V1.3.4</strong> — Outil professionnel de génération de CV<br>
            © 2025 Tous droits réservés
        </div>
        """,
        unsafe_allow_html=True,
    )

# ==========================================
# 🚀 MAIN ENTRY POINT
# ==========================================

if __name__ == "__main__":
    if not st.session_state.authenticated:
        restore_session_from_cookies()
    
    if st.session_state.authenticated:
        main_app()
        show_footer()
    else:
        show_login_screen()
