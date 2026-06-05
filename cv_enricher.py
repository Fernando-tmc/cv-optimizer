#!/usr/bin/env python3
"""
Universal CV Enricher
Reads any CV → Enriches with AI → Generates professional CV
"""

import os
import sys
import json
from docxtpl import DocxTemplate, RichText
from docx import Document
import jinja2
from typing import Dict, List, Any
import PyPDF2
import re
from zipfile import ZipFile
from xml.etree import ElementTree as ET

# === OCR IMPORTS ===
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import tempfile

print(">>> cv_enricher module loading", flush=True)


def fix_table_width_to_auto(doc):
    """
    Change table width from fixed to auto to prevent horizontal shift after merge.
    
    This fixes the issue where Skills Matrix tables with fixed width (e.g., 8.1 inches)
    get shifted right after merging because they don't fit within the page margins.
    
    Args:
        doc: Document object to fix
    
    Returns:
        int: Number of tables fixed
    """
    w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tables_fixed = 0
    
    for table in doc.tables:
        tbl = table._element
        tblPr = tbl.find(f'.//{w}tblPr')
        
        if tblPr is not None:
            # Find and fix tblW (table width)
            tblW = tblPr.find(f'.//{w}tblW')
            if tblW is not None:
                old_type = tblW.get(f'{w}type', 'unknown')
                old_w = tblW.get(f'{w}w', 'unknown')
                
                # Change to auto width
                tblW.set(f'{w}type', 'auto')
                tblW.set(f'{w}w', '0')
                
                print(f"   🔧 Table width changed: {old_type}={old_w} → auto=0")
                tables_fixed += 1
            
            # Remove fixed layout if present
            tblLayout = tblPr.find(f'.//{w}tblLayout')
            if tblLayout is not None:
                old_layout = tblLayout.get(f'{w}type', 'unknown')
                tblPr.remove(tblLayout)
                print(f"   🔧 Removed tblLayout: {old_layout}")
    
    return tables_fixed


class CVEnricher:
    """Universal CV enricher"""
    
    def __init__(self, api_key: str = None):
        """Initialize with Claude API key"""
        self.api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
        if not self.api_key:
            raise ValueError("❌ Claude API key missing! Set ANTHROPIC_API_KEY in Streamlit secrets or environment variable.")
        
        # Debug API key
        print(f">>> ANTHROPIC_KEY_PRESENT: {bool(self.api_key)}, len: {len(self.api_key) if self.api_key else 0}", flush=True)
        
        # Lazy loading
        self._anthropic_client = None
    
    def _get_anthropic_client(self):
        """Lazy loading of Anthropic client"""
        if self._anthropic_client is None:
            try:
                print(">>> Creating anthropic client", flush=True)
                import anthropic
                self._anthropic_client = anthropic.Anthropic(api_key=self.api_key)
                print(">>> Anthropic client created OK", flush=True)
            except Exception as e:
                print(f">>> ERROR creating anthropic client: {repr(e)}", flush=True)
                raise
        return self._anthropic_client
    
    # ========================================
    # MODULE 1 : EXTRACTION UNIVERSELLE
    # ========================================
    
    def detect_file_type(self, file_path: str) -> str:
        """Détecter le type de fichier"""
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            return 'pdf'
        elif ext in ['docx', 'doc']:
            return 'docx'
        elif ext in ['txt', 'text']:
            return 'txt'
        else:
            return 'unknown'
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extraire texte d'un PDF avec fallback OCR automatique
        1. Essaye PyPDF2 pour texte sélectionnable
        2. Si échec/texte vide → Utilise OCR sur images
        """
        print(f"📄 Extracting PDF: {file_path}", flush=True)
        
        try:
            # ===== ÉTAPE 1: Tentative extraction PyPDF2 =====
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                print(f"📊 PDF has {num_pages} pages", flush=True)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                    print(f"  Page {page_num}: {len(page_text) if page_text else 0} chars", flush=True)
            
            extracted_text = "\n".join(text).strip()
            
            # ===== VÉRIFIER SI L'EXTRACTION A FONCTIONNÉ =====
            # Seuil: Si moins de 100 caractères ou trop peu de mots → C'est scanné
            word_count = len(extracted_text.split())
            char_count = len(extracted_text)
            
            print(f"📈 PyPDF2 extraction: {char_count} chars, {word_count} words", flush=True)
            
            # Si extraction suffisante → Retourner
            if char_count > 100 and word_count > 20:
                print("✅ PDF text extraction successful (text-based PDF)", flush=True)
                return extracted_text
            
            # ===== ÉTAPE 2: PDF scanné détecté → OCR =====
            print("⚠️ PDF appears to be scanned (image-based). Switching to OCR...", flush=True)
            return self._extract_from_pdf_ocr(file_path)
            
        except Exception as e:
            print(f"❌ Error in PDF extraction: {e}", flush=True)
            # En cas d'erreur PyPDF2, essayer quand même OCR
            try:
                print("🔄 Trying OCR as fallback...", flush=True)
                return self._extract_from_pdf_ocr(file_path)
            except Exception as e2:
                print(f"❌ OCR fallback also failed: {e2}", flush=True)
                return ""
    
    def _extract_from_pdf_ocr(self, file_path: str) -> str:
        """
        Extraire texte d'un PDF scanné via OCR
        Utilise pdf2image + pytesseract
        """
        print("🔍 Starting OCR extraction...", flush=True)
        
        try:
            # Convertir PDF en images (une par page)
            # poppler_path peut être nécessaire sur Windows, mais pas sur Linux/Render
            images = convert_from_path(
                file_path,
                dpi=300,  # Haute résolution pour meilleur OCR
                fmt='jpeg',
                thread_count=2  # Parallélisation
            )
            
            print(f"📷 Converted {len(images)} pages to images", flush=True)
            
            # Extraire texte de chaque image
            all_text = []
            for i, image in enumerate(images, 1):
                print(f"  🔎 OCR processing page {i}/{len(images)}...", flush=True)
                
                # Appliquer OCR avec config optimisée
                # lang='eng+fra' pour anglais ET français
                page_text = pytesseract.image_to_string(
                    image,
                    lang='eng+fra',  # Anglais + Français
                    config='--psm 1 --oem 3'  # PSM 1 = automatic page segmentation with OSD
                )
                
                if page_text.strip():
                    all_text.append(f"--- Page {i} ---\n{page_text}")
                    print(f"  ✓ Page {i}: {len(page_text)} chars extracted", flush=True)
            
            extracted_text = "\n\n".join(all_text)
            print(f"✅ OCR extraction complete: {len(extracted_text)} chars total", flush=True)
            
            return extracted_text
            
        except Exception as e:
            print(f"❌ OCR extraction failed: {e}", flush=True)
            import traceback
            print(traceback.format_exc(), flush=True)
            return ""
    
     
    def extract_from_docx(self, file_path: str) -> str:
        """Extraire texte d'un Word + zones textes"""
        try:
            doc = Document(file_path)
            text = []
            
            # Paragraphes normaux
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())
            
            # Tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text.append(row_text)
            
            # NOUVEAU : Extraire les zones textes du XML
            textbox_content = self.extract_textboxes(file_path)
            if textbox_content:
                text.append("\n=== ZONES TEXTES ===")
                text.extend(textbox_content)
            
            return "\n".join(text)
        except Exception as e:
            print(f"⚠️ Erreur extraction Word: {e}")
            return ""
    def extract_from_txt(self, file_path: str) -> str:
        """Extraire texte d'un fichier texte"""
        try:
            # Essayer plusieurs encodages
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            # Si tout échoue, ignorer les erreurs
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"⚠️ Erreur extraction TXT: {e}")
            return ""
    
    def extract_textboxes(self, docx_path: str) -> list:
        """Extraire le contenu des zones textes (text boxes) du XML"""
        textboxes = []
        
        try:
            with ZipFile(docx_path, 'r') as docx:
                # Lire le document.xml
                xml_content = docx.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # Namespaces Word
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'v': 'urn:schemas-microsoft-com:vml',
                    'w10': 'urn:schemas-microsoft-com:office:word'
                }
                
                # Chercher tous les éléments de texte dans les zones textes
                # Les zones textes sont dans w:txbxContent
                for txbx in tree.findall('.//w:txbxContent', namespaces):
                    texts = []
                    for t in txbx.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text.strip())
                    if texts:
                        textboxes.append(' '.join(texts))
                
                # Aussi chercher dans v:textbox (ancien format)
                for vtxbx in tree.findall('.//v:textbox', namespaces):
                    texts = []
                    for t in vtxbx.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text.strip())
                    if texts:
                        textboxes.append(' '.join(texts))
                        
        except Exception as e:
            print(f"⚠️ Erreur extraction zones textes: {e}")
        
        return textboxes
    
    def extract_cv_text(self, cv_path: str) -> str:
        """Extraction universelle - détecte et extrait selon le type"""
        print(f"📄 Extraction du CV: {cv_path}")
        
        file_type = self.detect_file_type(cv_path)
        
        if file_type == 'pdf':
            print("   Format détecté: PDF")
            return self.extract_from_pdf(cv_path)
        elif file_type == 'docx':
            print("   Format détecté: Word")
            return self.extract_from_docx(cv_path)
        elif file_type == 'txt':
            print("   Format détecté: Texte")
            return self.extract_from_txt(cv_path)
        else:
            raise ValueError(f"❌ Format non supporté: {file_type}")

    # ========================================
    # MODULE 2 : PARSING INTELLIGENT
    # ========================================
    
    def parse_cv_with_claude(self, cv_text: str) -> Dict[str, Any]:
        """Parser le CV avec Claude pour extraire les infos structurées"""
        print("🤖 Parsing du CV avec Claude AI...", flush=True)
        
        try:
            client = self._get_anthropic_client()
            
            prompt = f"""Tu es un expert en analyse de CV. Extrait TOUTES les informations de ce CV et structure-les en JSON.

CV À ANALYSER:
{cv_text}

IMPORTANT CRITIQUE:
- Le NOM peut être caché dans un tableau HTML ou être stylisé. Cherche PARTOUT.
- Le LIEU DE RÉSIDENCE est OBLIGATOIRE : cherche "Montréal", "Montreal", villes + pays (ex: "Montreal CA", "Montréal, Canada", "Toronto ON", etc.). Si introuvable, mets "Location not specified".
- Les LANGUES sont OBLIGATOIRES : cherche "Français", "French", "English", "Anglais", "Bilingual", "Bilingue", etc. Si introuvable, mets ["Not specified"].

Extrait et structure en JSON STRICT (sans markdown):
{{
  "nom_complet": "Nom Prénom du candidat (cherche PARTOUT, même dans tableaux/HTML)",
  "titre_professionnel": "Titre/poste actuel",
  "profil_resume": "Résumé du profil si présent (sinon vide)",
  "lieu_residence": "OBLIGATOIRE - Ville, Pays (ex: Montréal, Canada) ou Montreal CA. Cherche codes pays (CA, US, FR). Si vraiment introuvable: 'Location not specified'",
  "langues": ["OBLIGATOIRE - Français", "Anglais", ... Cherche 'bilingual', 'French', 'English', etc. Si introuvable: ['Not specified']],
  "competences": ["compétence1", "compétence2", "compétence3", ...],
  "experiences": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre du poste",
      "responsabilites": ["tâche 1", "tâche 2", "tâche 3"]
    }}
  ],
  "formation": [
    {{
      "diplome": "Nom COMPLET du diplôme",
      "institution": "Nom école/université",
      "annee": "2020 (ou période exacte, sinon laisse vide)",
      "pays": "Pays SEULEMENT s'il est explicitement écrit dans le CV, sinon chaîne vide \"\". Ne JAMAIS inventer ni mettre Canada par défaut."
    }}
  ],
  "certifications": [
    {{
      "nom": "Nom certification",
      "organisme": "Organisme",
      "annee": "2023"
    }}
  ],
  "projets": [
    {{
      "nom": "Nom projet",
      "description": "Description courte"
    }}
  ]
}}

RÈGLES CRITIQUES:
- Le NOM est PRIORITAIRE - cherche dans tout le texte (tableaux, début, fin)
- LIEU DE RÉSIDENCE : cherche formats "Ville, Pays", "Montreal CA", "Montréal QC", codes postaux (H2X, etc.)
- LANGUES : cherche "Languages", "Langues", "French", "English", "Bilingual", même dans sections compétences
- Pour les diplômes: nom COMPLET + année EXACTE
- Extrait TOUT (ne rate rien)
- Si une section est vide, mets une liste vide []
- Format JSON strict uniquement"""

            print(f">>> Calling Claude API with timeout=300s...", flush=True)
            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                timeout=300.0,  # 5 minutes max
                messages=[{"role": "user", "content": prompt}]
            )
            print(f">>> API call completed successfully", flush=True)
            
        except Exception as e:
            print(f">>> ERROR calling anthropic for parsing: {repr(e)}", flush=True)
            return {}
        
        response_text = response.content[0].text.strip()
        
        # Nettoyer JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        try:
            parsed_data = json.loads(response_text)
            print(f"✅ Parsing réussi!")
            print(f"   Nom: [ANONYMIZED]")
            print(f"   Langues: {', '.join(parsed_data.get('langues', []))}")
            print(f"   Lieu: [ANONYMIZED]")
            print(f"   Compétences: {len(parsed_data.get('competences', []))}")
            print(f"   Expériences: {len(parsed_data.get('experiences', []))}")
            return parsed_data
        except json.JSONDecodeError as e:
            print(f"⚠️ Erreur JSON: {e}")
            print(f"Réponse brute: {response_text[:500]}")
            return {}

    # ========================================
    # MODULE 3 : ENRICHISSEMENT (TON PROMPT)
    # ========================================
    
    def read_job_description(self, jd_path: str) -> str:
        """Lire la job description"""
        file_type = self.detect_file_type(jd_path)
        
        if file_type == 'pdf':
            return self.extract_from_pdf(jd_path)
        elif file_type == 'docx':
            return self.extract_from_docx(jd_path)
        else:
            return self.extract_from_txt(jd_path)
    
    def analyze_cv_matching(self, parsed_cv: Dict[str, Any], jd_text: str, language: str = "French") -> Dict[str, Any]:
        """
        Analyser le matching entre CV et JD sans enrichir le contenu.
        Retourne uniquement: score_matching, domaines_analyses, synthese_matching
        """
        import time
        
        print(f"🔍 Analyse du matching CV/JD...", flush=True)
        
        start_time = time.time()
        
        try:
            client = self._get_anthropic_client()
            
            # Reconstruire le CV en texte pour le prompt
            cv_text = f"""
PROFIL: {parsed_cv.get('profil_resume', '')}

TITRE: {parsed_cv.get('titre_professionnel', '')}

COMPÉTENCES:
{chr(10).join(['- ' + comp for comp in parsed_cv.get('competences', [])])}

EXPÉRIENCES:
"""
            for exp in parsed_cv.get('experiences', []):
                cv_text += f"\n{exp.get('periode', '')} | {exp.get('entreprise', '')} | {exp.get('poste', '')}\n"
                for resp in exp.get('responsabilites', []):
                    cv_text += f"  - {resp}\n"
            
            cv_text += "\nFORMATION:\n"
            for form in parsed_cv.get('formation', []):
                cv_text += f"- {form.get('diplome', '')} | {form.get('institution', '')} | {form.get('annee', '')}\n"
        
            # PROMPT FOCALISÉ SUR L'ANALYSE DE MATCHING UNIQUEMENT - VERSION ULTRA-STRICTE V1.3.9
            prompt = f"""Tu es un système d'évaluation automatisé ULTRA-STRICT qui analyse le matching entre CV et Job Description.

🎯 ANALYSE DE MATCHING PONDÉRÉE (VERSION ULTRA-STRICTE V1.3.9):

⚠️ PRINCIPE FONDAMENTAL - ÉVALUATION ULTRA-RIGOUREUSE:
- Tu es un RECRUTEUR SENIOR EXTRÊMEMENT EXIGEANT avec 15+ ans d'expérience
- Tu recrutes pour des postes CRITIQUES où l'excellence est la norme
- CHAQUE point doit être MÉRITÉ avec des PREUVES CONCRÈTES du CV
- Si tu hésites entre 2 scores → TOUJOURS prends le PLUS BAS
- Agis comme si tu recrutais pour ton propre argent (zéro tolérance pour l'approximation)
- Pour le MÊME CV et la MÊME JD → EXACTEMENT le même score à chaque fois (cohérence algorithmique)

🔴 RÈGLE D'OR - SCORE GLOBAL = SOMME DOMAINES:
- Le score_matching FINAL = somme EXACTE de tous les scores de domaines
- VÉRIFIE 3 FOIS avant de répondre: somme des scores = score_matching
- Si tu calcules 58/100 en sommant les domaines → score_matching DOIT être 58
- NE JAMAIS inventer un score global différent de la somme calculée

═══════════════════════════════════════════════════
📋 ÉTAPE 1 - IDENTIFIER 5-8 DOMAINES CRITIQUES
═══════════════════════════════════════════════════

PROCESSUS AUTOMATIQUE D'IDENTIFICATION:
1. Scan complet de la JD - repérer TOUS les mots techniques/compétences
2. Compter la fréquence EXACTE de chaque technologie/compétence/méthodologie
3. Identifier les must-haves vs nice-to-haves
4. Créer une liste de domaines par ordre d'importance
5. Appliquer la formule de pondération ci-dessous

📊 FORMULE DE PONDÉRATION MATHÉMATIQUE:
Pour chaque domaine, calcule son poids avec:
Poids = (Mentions_JD × 10) + (Niveau_requis × 5) + Bonus_contexte

Où:
- Mentions_JD: Nombre de fois mentionné dans JD (1=once, 2=2-3 times, 3=4+ times)
- Niveau_requis: Must-have/Required=3, Important=2, Nice-to-have=1
- Bonus_contexte: +5 si dans le titre du poste, +3 si dans top requirements

💡 EXEMPLES DE DOMAINES TYPES:
- Technologies spécifiques (ex: "Python Django", "AWS Lambda", "React Native")
- Méthodologies (ex: "Agile/Scrum", "ITIL v4", "DevOps CI/CD")
- Compétences métier (ex: "Financial Modeling", "Clinical Trials Management")
- Certifications (ex: "PMP", "AWS Solutions Architect", "CPA")
- Langues avec niveau (ex: "Bilingual French/English C1+", "Spanish Business Level")
- Soft skills MESURABLES (ex: "Team Leadership 10+ people", "Stakeholder Management C-Suite")

⚠️ INTERDICTIONS ABSOLUES:
- NE JAMAIS créer de domaine vague type "General Fit", "Soft Skills", "Cultural Fit"
- NE JAMAIS créer de domaine "bonus" pour ajuster artificiellement le score
- TOUS les domaines doivent être EXPLICITEMENT mentionnés dans la JD
- Pas de domaines "catch-all" ou génériques

═══════════════════════════════════════════════════
🎯 ÉTAPE 2 - GRILLE D'ÉVALUATION ULTRA-STRICTE
═══════════════════════════════════════════════════

Pour CHAQUE domaine identifié, évalue le score avec cette GRILLE ULTRA-SÉVÈRE (0-100 points par domaine):

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔴 NIVEAU 0-15 POINTS: QUASI-AUCUNE COMPÉTENCE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
0 points: Compétence TOTALEMENT absente du CV (aucune mention directe ou indirecte)
10 points: Mention très vague OU compétence tangentielle (ex: "exposure to", "familiar with")
15 points: Mention superficielle OU formation théorique seulement SANS pratique OU <3 mois d'expérience

🟠 NIVEAU 20-35 POINTS: DÉBUTANT/JUNIOR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
20 points: 3-6 mois d'expérience pratique OU 1 projet simple réalisé sous supervision
25 points: 6-9 mois d'expérience OU 2 projets avec support d'équipe
30 points: 9-12 mois d'expérience avec autonomie partielle OU certification récente + pratique limitée
35 points: 1 an d'expérience solide avec quelques réalisations concrètes (mais sans metrics)

🟡 NIVEAU 40-55 POINTS: INTERMÉDIAIRE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
40 points: 1-1.5 ans d'expérience + 2-3 projets pertinents documentés
45 points: 1.5-2 ans d'expérience + contribution mesurable (ex: "improved X by Y%")
50 points: 2-2.5 ans d'expérience solide + réalisations quantifiées (metrics, budget, scope)
55 points: 2.5-3 ans + rôle de contributeur principal sur projets moyens

🟢 NIVEAU 60-75 POINTS: CONFIRMÉ/SENIOR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
60 points: 3-4 ans d'expérience confirmée + ownership de projets + résultats mesurables
65 points: 4-5 ans + expertise démontrée par réalisations significatives (ex: led team of 5, managed $500K budget)
70 points: 5-6 ans + rôle de lead/expert technique + mentorship + process improvements
75 points: 6-7 ans + expertise reconnue EN INTERNE (promotions, leadership technique, formations données en interne)

🔵 NIVEAU 80-90 POINTS: EXPERT EXCEPTIONNEL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
80 points: 7-8 ans d'expérience TRÈS solide + leadership prouvé + expertise reconnue PAR L'INDUSTRIE (speaking engagements, certifications avancées, articles techniques)
85 points: 8-10 ans + contribution MAJEURE à l'industrie (architecture de solutions complexes multi-millions, thought leadership, certifications rares)
90 points: 10-12 ans + expertise de NIVEAU MONDIAL dans ce domaine spécifique (publications académiques/industrie, conférences internationales, mentor d'experts, awards/recognition)

🏆 NIVEAU 95-100 POINTS: QUASI-IMPOSSIBLE - TOP 0.1% MONDIAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
95 points: 12-15 ans + reconnaissance INTERNATIONALE + contributions MAJEURES à l'évolution du domaine (patents, standards, books, keynote speaker top conferences)
100 points: RÉSERVÉ AUX LÉGENDES VIVANTES - 15+ ans + autorité MONDIALE incontestée dans le domaine + impact transformationnel sur l'industrie (ex: créateur de framework utilisé par millions, membre de comités internationaux, consultant pour Fortune 10)

⚠️ RÈGLES ULTRA-STRICTES D'ATTRIBUTION:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. JAMAIS de score ≥60 sans PREUVES QUANTIFIÉES concrètes dans le CV
2. JAMAIS de score ≥75 sans leadership/mentorship/expertise reconnue PROUVÉE
3. JAMAIS de score ≥85 sans contributions MAJEURES à l'industrie (publications, speaking, thought leadership)
4. JAMAIS de score ≥95 sans reconnaissance INTERNATIONALE vérifiable
5. Si le CV mentionne l'expérience en années SEULEMENT sans détails de réalisations → score MAX = 55
6. Si aucun metric/chiffre fourni pour un domaine → score MAX = 50
7. Si le candidat change de domaine/technologie fréquemment (job hopping) → pénalité de -10 points
8. Certifications SANS expérience pratique associée → score MAX = 30
9. Expérience dans environnement non-professionnel (side projects, freelance) compte pour 50% seulement
10. Si tu hésites entre 2 scores → TOUJOURS choisir le PLUS BAS

⚙️ RÈGLES DE CALCUL FINAL:
1. Score brut du domaine = évaluation selon grille ci-dessus (0-100)
2. Score pondéré = (score_brut × poids) / 100
3. Score_max du domaine = poids

Exemple détaillé:
- Domaine: "Python Backend Development" | Poids: 25%
- Candidat: 4.5 ans d'expérience Python, 3 projets documentés, led team of 3, aucune publication
- Évaluation: Entre 60 et 65 points → choisir 60 (règle du plus bas)
- Score pondéré: (60 × 25) / 100 = 15 points
- Score_max: 25 points
- Notation: 15/25

═══════════════════════════════════════════════════
📊 ÉTAPE 3 - CALCULER LE SCORE TOTAL
═══════════════════════════════════════════════════

Score_matching = SOMME de tous les scores pondérés (arrondi à l'entier)

Exemple:
15 (Python) + 10 (AWS) + 8 (Agile) + 12 (API Design) + 9 (PostgreSQL) + 7 (Docker) = 61/100

⚠️ VÉRIFICATIONS FINALES OBLIGATOIRES:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 🔴 CRITIQUE: Somme des poids = EXACTEMENT 100% (PAS 99%, PAS 101%, PAS 110%, EXACTEMENT 100%)
   - Additionne TOUS les "poids" avant de répondre
   - Si total ≠ 100% → AJUSTE les poids proportionnellement pour totaliser exactement 100%
   - Exemple: Si tu as 110%, divise chaque poids par 1.1 (25%→22.7%, 20%→18.2%, etc.)
   - Vérifie 2 fois: somme finale des poids DOIT être 100
2. Score_matching = somme EXACTE des scores pondérés
3. Si score > 80 → TRIPLE-CHECK: y a-t-il vraiment des preuves d'expertise exceptionnelle?
4. Si score > 90 → QUADRUPLE-CHECK: est-ce vraiment un candidat top 1% mondial? (la réponse devrait presque toujours être NON)
5. Refaire le calcul 2 fois pour confirmer

🎯 PHILOSOPHIE DE NOTATION ATTENDUE (distribution réaliste):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- Score 95-100: <1% des candidats (quasi-impossible, réservé aux légendes)
- Score 85-94: ~5% (top performers exceptionnels)
- Score 75-84: ~15% (très bons candidats confirmés)
- Score 65-74: ~25% (bons candidats solides)
- Score 50-64: ~30% (candidats acceptables avec gaps)
- Score <50: ~24% (candidats insuffisants)

⚠️ DERNIÈRE VÉRIFICATION AVANT RÉPONSE:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Pose-toi ces questions pour CHAQUE domaine où tu as donné ≥60 points:
- Ai-je des PREUVES CONCRÈTES d'expérience quantifiable dans le CV?
- Ai-je des RÉALISATIONS MESURABLES (metrics, budget, team size, impact)?
- Le candidat a-t-il eu un rôle de LEADERSHIP/OWNERSHIP démontré?
- Pour les scores ≥85: y a-t-il des contributions à l'INDUSTRIE (publications, speaking, thought leadership)?
Si la réponse n'est pas un OUI catégorique avec preuves multiples → BAISSE le score.

═══════════════════════════════════════════════════
📝 ÉTAPE 4 - SYNTHÈSE EXECUTIVE (4-5 LIGNES MAX)
═══════════════════════════════════════════════════

Rédige une synthèse ULTRA-CONCISE en 4-5 LIGNES (80-100 mots maximum) qui:

STRUCTURE OBLIGATOIRE (1 paragraphe fluide):
1. Lead with match level + score (e.g., "GOOD match (73/100) for [Role]")
2. Highlight 2-3 TOP strengths with brief evidence (years, key achievement, metric)
3. Mention 1-2 minor gaps or "nice-to-haves" missing
4. End with clear recommendation: "Interview - [reason]" or "Pass - [reason]"

EXEMPLE FORMAT:
"GOOD match (73/100) for Senior Full-Stack Developer. Strong Python backend (8 years) with proven cloud migration leadership (60% deployment time reduction). Full-stack capability confirmed with React + modern DevOps. Minor gaps: Kubernetes nice-to-have, limited Montreal-specific experience. Recommendation: Interview - solid technical fit with measurable impact."

RÈGLES CRITIQUES:
- MAX 4-5 lignes (80-100 mots)
- NO paragraphs, NO bullet points - juste 1 bloc de texte fluide
- Include score + match level (EXCELLENT 85+, GOOD 70-84, MODERATE 55-69, WEAK <55)
- Be specific with numbers/metrics when available
- Professional but direct tone
- Clear go/no-go recommendation at the end

═══════════════════════════════════════════════════
📄 FORMAT DE SORTIE JSON
═══════════════════════════════════════════════════

📄 JOB DESCRIPTION:
{jd_text}

📄 CV DU CANDIDAT:
{cv_text}

═══════════════════════════════════════════════════

🎯 GÉNÈRE MAINTENANT TON ANALYSE - FORMAT JSON STRICT:

Retourne UNIQUEMENT un JSON avec cette structure (sans texte avant/après):

{{
    "score_matching": 58,
    "domaines_analyses": [
        {{
            "domaine": "Nom du domaine technique/compétence exact",
            "poids": 20,
            "score": 10,
            "score_max": 20,
            "match": "bon",
            "commentaire": "Justification FACTUELLE ultra-détaillée basée sur des éléments PRÉCIS du CV avec années d'expérience, projets, réalisations, metrics. Minimum 2-3 phrases complètes."
        }}
    ],
    "synthese_matching": "COMPREHENSIVE PROFESSIONAL ANALYSIS (4-6 DETAILED PARAGRAPHS, 250-350 WORDS):

[Paragraph 1 - Overall Assessment]
[Detailed assessment text...]

[Paragraph 2 - Top Strengths]
[Detailed strengths text...]

[Paragraph 3 - Partial Matches]
[Detailed partial matches text...]

[Paragraph 4 - Gaps]
[Detailed gaps text...]

[Paragraph 5 - Final Recommendation]
[Detailed recommendation text...]"
}}

⚠️ RÈGLES JSON CRITIQUES:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- "match" peut être: "excellent" (≥85/100), "bon" (65-84), "partiel" (40-64), "incompatible" (<40)
- Tous les scores doivent être des NOMBRES (pas de strings)
- 🔴 La somme des poids doit faire EXACTEMENT 100 (vérifie 2 fois avant de répondre)
- Le score_matching doit être la somme exacte des scores de tous les domaines
- Commentaire: minimum 2-3 phrases complètes avec détails factuels précis du CV
- Synthèse: MAXIMUM 4-5 lignes (80-100 mots), format executive summary

⚠️ LANGUE DE SORTIE: TOUT le texte (noms de domaines, commentaires, synthèse) doit être rédigé en {language}.
- Si {language} = "French": noms de domaines, commentaires et synthèse 100% en français.
- Si {language} = "English": domain names, comments and synthesis 100% in English.

Génère l'analyse maintenant:"""
            
            print(f">>> Calling Claude API for matching analysis...", flush=True)
            
            # ✅ RETRY LOGIC FOR TIMEOUTS
            max_retries = 2
            response = None
            last_error = None
            
            for attempt in range(max_retries):
                try:
                    response = client.messages.create(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=4000,
                        timeout=900.0,  # 15 minutes
                        messages=[{"role": "user", "content": prompt}]
                    )
                    break  # Success - exit retry loop
                    
                except Exception as e:
                    last_error = e
                    error_name = type(e).__name__
                    
                    # Check if it's a timeout error
                    if 'timeout' in error_name.lower() or 'timeout' in str(e).lower():
                        if attempt < max_retries - 1:
                            print(f"⏱️ Timeout attempt {attempt+1}/{max_retries}, retrying...", flush=True)
                            continue
                        else:
                            # Final timeout - return error result
                            print(f"❌ Final timeout after {max_retries} attempts", flush=True)
                            return {
                                'error': 'timeout',
                                'score_matching': 0,
                                'domaines_analyses': [],
                                'synthese_matching': "⏱️ L'analyse a pris trop de temps (timeout après plusieurs tentatives). Veuillez réessayer avec un CV plus court ou contactez le support."
                            }
                    else:
                        # Other error - re-raise
                        raise
            
            if response is None:
                # Should not happen, but safety check
                raise last_error
            
            # Extraire tokens
            usage = response.usage
            input_tokens = usage.input_tokens
            output_tokens = usage.output_tokens
            total_tokens = input_tokens + output_tokens
            
            print(f">>> API Response received. Tokens: {total_tokens}", flush=True)
            
            # Parser la réponse
            response_text = response.content[0].text.strip()
            
            # Nettoyer le JSON
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()
            
            # Parser le JSON
            try:
                matching_result = json.loads(response_text)
                print(f">>> JSON parsed successfully!", flush=True)
                
                # V1.3.5 FIX ULTIME: Recalculer TOUS les scores pondérés pour garantir cohérence
                if 'domaines_analyses' in matching_result and matching_result['domaines_analyses']:
                    # Vérifier si Claude a mis les scores BRUTS (0-100) au lieu des scores pondérés
                    # Indice: Si la somme des scores > 100, ce sont des scores bruts
                    
                    total_weight = sum(d.get('poids', 0) for d in matching_result['domaines_analyses'])
                    sum_scores = sum(d.get('score', 0) for d in matching_result['domaines_analyses'])
                    
                    # Si somme des scores > 100 OU > total_weight → Ce sont des scores BRUTS, il faut recalculer
                    if sum_scores > total_weight:
                        print(f"⚠️ Scores bruts détectés (somme={sum_scores}) → Recalcul des scores pondérés")
                        
                        # Recalculer chaque score pondéré: (score_brut × poids) / 100
                        for domain in matching_result['domaines_analyses']:
                            score_brut = domain.get('score', 0)
                            poids = domain.get('poids', 0)
                            # Le score doit être le score pondéré, pas le brut
                            score_pondere = (score_brut * poids) / 100
                            domain['score'] = round(score_pondere)
                            domain['score_max'] = poids
                            print(f"   {domain['domaine'][:40]}: {score_brut}/100 × {poids}% = {round(score_pondere)}/{poids}")
                        
                        # Recalculer le total
                        calculated_score = sum(d.get('score', 0) for d in matching_result['domaines_analyses'])
                    else:
                        # Les scores sont déjà pondérés
                        calculated_score = sum_scores
                    
                    # Normaliser si les poids dépassent 100%
                    if total_weight > 100:
                        print(f"⚠️ Poids totaux: {total_weight}% → Normalisation à 100%")
                        calculated_score = (calculated_score / total_weight) * 100
                    
                    original_score = matching_result.get('score_matching', 0)
                    
                    # Utiliser le score calculé (toujours plus fiable)
                    final_score = min(round(calculated_score), 100)
                    
                    if abs(final_score - original_score) > 2:
                        print(f"⚠️ Score mismatch: Claude={original_score}, Calculated={final_score}")
                        print(f"   Using calculated score: {final_score}/100")
                    
                    matching_result['score_matching'] = final_score
                    
                    # ✅ Update synthese_matching with correct score if needed
                    if abs(final_score - original_score) > 2 and 'synthese_matching' in matching_result:
                        synthese = matching_result['synthese_matching']
                        # Replace score mentions in common formats
                        import re
                        # Format: "score of XX" or "XX/100" or "XX out of 100"
                        synthese = re.sub(
                            rf'\b{original_score}/100\b',
                            f'{matching_result["score_matching"]}/100',
                            synthese
                        )
                        synthese = re.sub(
                            rf'\bscore of {original_score}\b',
                            f'score of {matching_result["score_matching"]}',
                            synthese,
                            flags=re.IGNORECASE
                        )
                        synthese = re.sub(
                            rf'\b{original_score} out of 100\b',
                            f'{matching_result["score_matching"]} out of 100',
                            synthese,
                            flags=re.IGNORECASE
                        )
                        matching_result['synthese_matching'] = synthese
                        print(f"   ✅ Updated synthese_matching to reflect corrected score: {matching_result['score_matching']}/100")
                        
            except json.JSONDecodeError as e:
                print(f"⚠️ JSON Error: {e}", flush=True)
                print(f">>> Attempting to fix JSON...", flush=True)
                
                # Tentative de réparation
                fix_prompt = f"""The following JSON is malformed. Please fix it and return ONLY the corrected JSON without any explanation or markdown:

{response_text}

Return the corrected JSON directly:"""
                
                # ✅ RETRY LOGIC FOR JSON FIX
                fix_response = None
                for fix_attempt in range(2):
                    try:
                        fix_response = client.messages.create(
                            model="claude-sonnet-4-5-20250929",
                            max_tokens=4000,
                            timeout=300.0,  # 5 minutes for fix
                            messages=[{"role": "user", "content": fix_prompt}]
                        )
                        break
                    except Exception as fix_error:
                        if 'timeout' in type(fix_error).__name__.lower() or 'timeout' in str(fix_error).lower():
                            if fix_attempt < 1:
                                print(f"⏱️ JSON fix timeout, retrying...", flush=True)
                                continue
                            else:
                                # Can't fix JSON - return error
                                return {
                                    'error': 'json_parse_timeout',
                                    'score_matching': 0,
                                    'domaines_analyses': [],
                                    'synthese_matching': "❌ Erreur de parsing JSON et timeout lors de la correction. Veuillez réessayer."
                                }
                        else:
                            raise
                
                if fix_response is None:
                    return {
                        'error': 'json_fix_failed',
                        'score_matching': 0,
                        'domaines_analyses': [],
                        'synthese_matching': "❌ Impossible de corriger le JSON malformé."
                    }
                
                fixed_text = fix_response.content[0].text.strip()
                if fixed_text.startswith('```json'):
                    fixed_text = fixed_text[7:]
                if fixed_text.startswith('```'):
                    fixed_text = fixed_text[3:]
                if fixed_text.endswith('```'):
                    fixed_text = fixed_text[:-3]
                fixed_text = fixed_text.strip()
                
                matching_result = json.loads(fixed_text)
                print(f">>> JSON successfully fixed and parsed!", flush=True)
            
            # Calculer le temps et coût
            processing_time = round(time.time() - start_time, 2)
            cost_input = (input_tokens / 1_000_000) * 3.0
            cost_output = (output_tokens / 1_000_000) * 15.0
            total_cost = round(cost_input + cost_output, 4)
            
            # Ajouter les métadonnées
            matching_result['_metadata'] = {
                'processing_time_seconds': processing_time,
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens,
                'estimated_cost_usd': total_cost
            }
            
            print(f"✅ Analyse de matching réussie!")
            print(f"   Score matching: {matching_result.get('score_matching', 0)}/100")
            print(f"   Domaines analysés: {len(matching_result.get('domaines_analyses', []))}")
            print(f"   ⏱️ Temps: {processing_time}s")
            print(f"   📊 Tokens: {total_tokens:,}")
            print(f"   💰 Coût: ${total_cost}")
            
            return matching_result
            
        except Exception as e:
            print(f"❌ Erreur analyse matching: {e}", flush=True)
            import traceback
            print(traceback.format_exc(), flush=True)
            return {
                'score_matching': 0,
                'domaines_analyses': [],
                'synthese_matching': f'Erreur lors de l\'analyse: {str(e)}'
            }
    
    def enrich_cv_with_prompt(
        self, 
        parsed_cv: Dict[str, Any], 
        jd_text: str, 
        language: str = "French",
        matching_analysis: Dict[str, Any] = None  # ✅ FIX: Nouveau paramètre pour réutiliser le matching
    ) -> Dict[str, Any]:
        """
        Enrichir le CV avec l'IA
        
        Args:
            parsed_cv: CV parsé
            jd_text: Job Description
            language: Langue cible (French/English)
            matching_analysis: Résultat optionnel du matching préalable (Step 1)
                              Si fourni, réutilise le score au lieu de le recalculer
        
        Returns:
            CV enrichi avec tous les champs nécessaires
        """
        import time
        
        # ⚠️ CRITICIAL: Déterminer si on réutilise le scoring du Step 1
        reuse_scoring = matching_analysis is not None
        
        print(f"✨ Enrichissement du CV avec l'IA...", flush=True)
        print(f"   Langue cible: {language}", flush=True)
        print(f"   Mode: {'Réutilisation scoring Step 1' if reuse_scoring else 'Scoring complet'}", flush=True)
        
        # ⏱️ Démarrer le chronomètre
        start_time = time.time()
        
        try:
            client = self._get_anthropic_client()
            
            # Reconstruire le CV en texte pour le prompt
            cv_text = f"""
PROFIL: {parsed_cv.get('profil_resume', '')}

TITRE: {parsed_cv.get('titre_professionnel', '')}

COMPÉTENCES:
{chr(10).join(['- ' + comp for comp in parsed_cv.get('competences', [])])}

EXPÉRIENCES:
"""
            for exp in parsed_cv.get('experiences', []):
                cv_text += f"\n{exp.get('periode', '')} | {exp.get('entreprise', '')} | {exp.get('poste', '')}\n"
                for resp in exp.get('responsabilites', []):
                    cv_text += f"  - {resp}\n"
            
            cv_text += "\nFORMATION:\n"
            for form in parsed_cv.get('formation', []):
                cv_text += f"- {form.get('diplome', '')} | {form.get('institution', '')} | {form.get('annee', '')}\n"
            
            # ✅ FIX: Ajouter PROJETS pour traduction
            cv_text += "\nPROJETS PERTINENTS:\n"
            for projet in parsed_cv.get('projets', []):
                cv_text += f"- {projet}\n"
            
            # ✅ FIX: Ajouter CERTIFICATIONS pour traduction
            cv_text += "\nCERTIFICATIONS:\n"
            for cert in parsed_cv.get('certifications', []):
                cert_name = cert.get('nom', cert.get('name', ''))
                cert_org = cert.get('organisme', cert.get('institution', ''))
                cert_year = cert.get('annee', cert.get('year', ''))
                cv_text += f"- {cert_name} | {cert_org} | {cert_year}\n"
        
            # PROMPT ULTRA-RENFORCÉ POUR COHÉRENCE ABSOLUE
            language_instruction = f"""
🚨 RÈGLE ABSOLUE - LANGUE {language.upper()} 🚨
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ INSTRUCTION CRITIQUE - LANGUE OBLIGATOIRE: {language.upper()}

Tu DOIS générer 100% du contenu en {language} - AUCUNE EXCEPTION:
✓ Le TITRE PROFESSIONNEL en {language}
✓ Le PROFIL ENRICHI en {language}
✓ TOUTES les COMPÉTENCES en {language}
✓ TOUTES les EXPÉRIENCES en {language}
✓ TOUS les noms de catégories en {language}
✓ TOUTES les descriptions en {language}
✓ TOUS les mots-clés en {language}

🔴 SI {language} = "French":
- Utilise: "Analyste", "Gestion", "Configuration", "Développement", "Senior"
- PAS: "Analyst", "Management", "Development"
- Exemple titre: "Analyste QA Senior" ✓ (PAS "Senior QA Analyst" ✗)
- Exemple description: "Configuration de SharePoint incluant gestion..."
- Style: Français professionnel standard

🔴 SI {language} = "English":
- Utilise: "Analyst", "Management", "Configuration", "Development", "Senior"
- PAS: "Analyste", "Gestion", "Développement"
- Exemple titre: "Senior QA Analyst" ✓ (PAS "Analyste QA Senior" ✗)
- Exemple description: "SharePoint configuration including management..."
- Style: Professional English standard

IMPORTANT TITRE PROFESSIONNEL:
- Adapte le titre à la Job Description
- Le titre doit être COURT (3-5 mots maximum)
- Le titre doit être en {language} - VÉRIFIE 2 FOIS
- Si langue = French: ordre français (ex: "Analyste Configuration SharePoint")
- Si langue = English: ordre anglais (ex: "SharePoint Configuration Analyst")

VÉRIFICATION FINALE OBLIGATOIRE:
Avant de répondre, relis TOUT ton JSON et confirme que:
1. Le titre_professionnel_enrichi est en {language} ✓
2. Le profil_enrichi est en {language} ✓
3. Toutes les catégories de compétences sont en {language} ✓
4. Toutes les descriptions sont en {language} ✓
5. Les responsabilités des expériences sont en {language} ✓

Si UNE SEULE phrase n'est pas en {language} → RECOMMENCE TOUT.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🎯 RÔLE CRITIQUE - TU ES UN RECRUTEUR SENIOR PROFESSIONNEL:
- Tu as 15+ ans d'expérience en recrutement technique
- Tu travailles pour le CLIENT (l'entreprise qui recrute)
- Ta mission: évaluer si le CANDIDAT correspond EXACTEMENT aux besoins du CLIENT
- Tu dois être OBJECTIF, RIGOUREUX et REPRODUCTIBLE dans ton évaluation
- Ton scoring doit être IDENTIQUE si tu analyses le même CV/JD plusieurs fois
- Tu notes comme un examinateur professionnel, pas comme un vendeur
"""
            
            # ✅ FIX: Choisir le prompt selon si on réutilise le matching ou non
            if reuse_scoring:
                # ============================================
                # VERSION SIMPLIFIÉE - Matching déjà fait au Step 1
                # ============================================
                prompt = f"""Voici la job description et le CV actuel ci-dessous.

🔹 Reformate et reformule LÉGÈREMENT le CV ci-dessous au format TMC, en gardant un ton professionnel sobre.
{language_instruction}

🚫🚫 RÈGLE N°1 — FIDÉLITÉ ABSOLUE, ZÉRO INVENTION (prime sur tout le reste) 🚫🚫
- Utilise UNIQUEMENT des informations RÉELLEMENT présentes dans le CV source ci-dessous.
- INTERDIT d'ajouter une compétence, technologie, outil, chiffre, résultat ou responsabilité absent du CV.
- INTERDIT de "compléter" avec des mots-clés de la Job Description que le candidat n'a pas. La JD sert SEULEMENT à choisir quels éléments VRAIS du CV mettre en avant et dans quel ordre — JAMAIS à ajouter ce qui est absent.
- Si une info n'est pas dans le CV (résultat chiffré, outil...), NE L'INVENTE PAS.
- En cas de doute : ne le mets pas.

🎙️ TON — sobre et factuel (surtout PAS "IA"/marketing) :
- Style professionnel neutre, comme un consultant le rédigerait.
- INTERDIT : superlatifs/formules marketing ("exceptionnel", "expert reconnu", "passionné", "maîtrise parfaite", "solide expérience", "dynamique"...).
- Phrases simples et directes, sans emphase artificielle.

⚠️ L'analyse de matching est DÉJÀ faite. Tu fais UNIQUEMENT la reformulation fidèle du contenu.

Fais :

1. TITRE court en {language} (3-5 mots max), basé sur le VRAI métier du candidat (orienté JD seulement si ça reste vrai).

2. PROFIL: paragraphe factuel de 4-5 lignes en {language}, reformulant le profil RÉEL (pas de superlatifs). 3-5 technologies réellement maîtrisées en **gras**.

3. EXPÉRIENCES — reformulation LÉGÈRE (jamais de copié-collé mot à mot, jamais d'invention) :
   - Style NOMINAL uniforme : chaque ligne commence par un nom d'action ("Conception de…", "Réalisation de…", "Développement de…", "Migration de…", "Analyse de…").
   - JAMAIS à la 1re personne ("j'ai…"), JAMAIS à l'infinitif en tête ("Réaliser…").
   - Mêmes FAITS que le CV source, juste mieux formulés et homogènes.
   - Bullets courts (1 ligne), 4-6 par expérience max.

Réponds en JSON STRICT (sans markdown) avec cette structure:
{{
  "titre_professionnel_enrichi": "TITRE COURT en {language} (3-5 mots max)",
  
  "profil_enrichi": "Profil NARRATIF 5-6 lignes en {language} avec **3-5 technologies clés** en gras",
  
  "mots_cles_a_mettre_en_gras": ["Technologies RÉELLEMENT présentes dans le CV (recoupant éventuellement la JD) - PAS de technos absentes du CV"],
  
  "competences_enrichies": {{
    "Nom Catégorie 1 (3-6 mots max)": [
      "**Compétence réelle** : description courte (100-150 caractères max) UNIQUEMENT à partir de ce qui figure dans le CV. N'ajoute aucun outil ni résultat absent du CV.",
      "**Autre compétence réelle** : description factuelle et concise, sans invention."
    ],
    "Nom Catégorie 2": [
      "Compétence concise..."
    ]
  }},
  
  RÈGLES ULTRA-CRITIQUES pour les compétences (NON-NÉGOCIABLE):
  - Noms de catégories COURTS (3-6 mots max)
  - 5-6 catégories ADAPTÉES à la JD
  - Chaque catégorie: 3-5 compétences MAXIMUM
  - CHAQUE compétence : 2-3 LIGNES MAXIMUM (100-150 caractères) - NE PAS DÉPASSER
  - N'inclus QUE des compétences réellement présentes dans le CV. AUCUNE compétence inventée, même si la JD la demande.
  - Descriptions factuelles et sobres, sans superlatifs ni résultats chiffrés inventés
  - 1-3 technologies RÉELLES en **gras** par compétence
  
  "experiences_enrichies": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre RÉEL du poste (reformulé seulement si fidèle)",
      "responsabilites": [
        "Configuration **Open edX** incluant structuration et intégration avec **SharePoint** pour gestion contenus",
        "Automatisation processus documentaires via **Power Automate** et **Teams** pour améliorer efficacité"
      ],
      "environment": "**Open edX**, **SharePoint**, **Microsoft 365**, Teams, Power Automate, OneDrive, SQL"
    }}
  ],
  
  "formation_enrichie": [
    {{
      "institution": "Institution name in {language}",
      "diplome": "Degree title in {language}",
      "annee": "2018"
    }}
  ],
  
  "projets_enrichis": [
    "Project description 1 in {language}",
    "Project description 2 in {language}"
  ]
}}

FORMAT OBLIGATOIRE (COPIER format compétences):
- Responsabilités: Technologies **isolées** dans texte normal (ex: "Configuration **Tech1** incluant **Tech2** pour résultats")
- Environnement: Liste virgules avec 3-5 technologies **critiques** en gras, autres sans
- JAMAIS phrases entières en gras
- Maximum 2-3 mots entre **astérisques**

---

JOB DESCRIPTION:
{jd_text}

---

CV ACTUEL:
{cv_text}

---

IMPORTANT FINAL - RÈGLES JSON STRICTES:
- Génère UNIQUEMENT du JSON valide
- PAS de commentaires (// ou /* */)
- PAS de virgules finales (trailing commas)
- PAS de markdown (```json ou ```)
- TOUS les strings doivent utiliser des guillemets doubles ""
- Vérifie que TOUTES les accolades et crochets sont fermés
- Si tu hésites sur un champ, mets une valeur par défaut plutôt qu'une erreur

Réponds UNIQUEMENT avec du JSON pur, sans rien d'autre avant ou après."""

            else:
                # ============================================
                # VERSION COMPLÈTE - Mode legacy/fallback avec matching inclus
                # ============================================
                prompt = f"""Voici la job description et le CV actuel ci-dessous.

🔹 Reformate et reformule LÉGÈREMENT le CV au format TMC, ton professionnel sobre.
{language_instruction}

🚫🚫 RÈGLE N°1 — FIDÉLITÉ ABSOLUE, ZÉRO INVENTION (prime sur tout) 🚫🚫
- Utilise UNIQUEMENT ce qui est RÉELLEMENT dans le CV source. INTERDIT d'ajouter compétence, techno, outil, chiffre ou résultat absent du CV.
- La JD sert à choisir quels éléments VRAIS mettre en avant, JAMAIS à ajouter ce que le candidat n'a pas.
- Ton sobre/factuel : INTERDIT les superlatifs et formules marketing ("exceptionnel", "expert reconnu", "passionné"...).
- Expériences : style NOMINAL ("Conception de…", "Réalisation de…"), jamais à la 1re personne, jamais d'invention.

🎯 ANALYSE DE MATCHING PONDÉRÉE (ULTRA-CRITIQUE - COHÉRENCE ABSOLUE REQUISE):

⚠️ PRINCIPE FONDAMENTAL DE COHÉRENCE - MÉTHODOLOGIE STRICTE:
- Tu es un SYSTÈME D'ÉVALUATION AUTOMATISÉ, pas un humain
- Pour le MÊME CV et la MÊME JD → EXACTEMENT le même score à chaque fois
- Utilise une grille d'évaluation MATHÉMATIQUE et REPRODUCTIBLE
- Agis comme un ALGORITHME, pas comme un recruteur subjectif
- Chaque critère suit des règles BINAIRES strictes (oui/non, présent/absent)
- Tu DOIS pouvoir justifier CHAQUE point attribué avec des FAITS du CV
- Si tu hésites entre 2 scores → prends le PLUS BAS (principe de strictness)

🔴 RÈGLE D'OR - SCORE GLOBAL = SOMME DOMAINES:
- Le score_matching FINAL = somme EXACTE de tous les scores de domaines
- VÉRIFIE 3 FOIS avant de répondre: somme des scores = score_matching
- Si tu calcules 37/100 en sommant les domaines → score_matching DOIT être 37
- NE JAMAIS inventer un score global différent de la somme calculée

ÉTAPE 1 - IDENTIFIER 5-8 DOMAINES CRITIQUES (MÉTHODE ALGORITHIMQUE):

📋 PROCESSUS AUTOMATIQUE D'IDENTIFICATION:
1. Scan complet de la JD - repérer TOUS les mots techniques
2. Compter la fréquence EXACTE de chaque technologie/compétence
3. Créer une liste de domaines par ordre d'importance
4. Appliquer la formule de pondération ci-dessous

📊 FORMULE DE PONDÉRATION MATHÉMATIQUE:
Pour chaque domaine, calcule son poids avec:

Poids_Base = (Nombre_mentions / Total_mentions_techniques) × 100

Bonus:
- +20% si c'est le TITRE du poste (ex: ".NET Developer" → Stack .NET = +20%)
- +15% si mots "Required", "Must have", "Essential", "Critical"
- +10% si mentionné dans les 3 premières lignes de la JD
- +5% par occurrence au-delà de 3 mentions

Poids_Final = min(Poids_Base + Bonus, 50%)  ← Aucun domaine ne peut dépasser 50%

RÈGLES STRICTES DE PONDÉRATION:
- Stack technique principal (dans titre ou 5+ mentions): 30-50%
- Compétences techniques secondaires (3-4 mentions): 15-25%
- Compétences techniques tertiaires (1-2 mentions): 5-15%
- Soft skills/Leadership: MAXIMUM 10% (sauf si poste management)
- TOTAL des poids = EXACTEMENT 100% (vérifie avec calculatrice)
- Si total ≠ 100%, ajuste proportionnellement tous les poids

ÉTAPE 2 - SCORER CHAQUE DOMAINE (ALGORITHME DE NOTATION STRICT):

🤖 SYSTÈME DE NOTATION AUTOMATISÉ - APPLIQUE CES RÈGLES EXACTEMENT:

POUR CHAQUE DOMAINE, COMPTE:
a) Nombre de mentions EXACTES de la technologie dans le CV
b) Nombre de projets/expériences utilisant cette technologie  
c) Durée totale d'utilisation (années)
d) Niveau démontré (junior/intermédiaire/senior)

📐 FORMULE MATHÉMATIQUE DE SCORING:

Étape 2A - Score Brut (0-100%):
• 0% : ZÉRO mention de la techno dans le CV, stack incompatible
• 10% : Technologie proche mentionnée (PostgreSQL pour SQL Server)
• 25% : 1 mention + aucune expérience pratique (formation seulement)
• 40% : 1-2 mentions + 1 projet + <1 an d'expérience
• 60% : 3-4 mentions + 2 projets + 1-2 ans d'expérience
• 80% : 5+ mentions + 3+ projets + 3+ ans d'expérience
• 100% : 7+ mentions + expertise démontrée + senior confirmé

Étape 2B - Ajustements OBLIGATOIRES:
• Si stack incompatible (Java vs .NET) → Score = 0% (NON-NÉGOCIABLE)
• Si technologie absente du CV → Score = 0% (NON-NÉGOCIABLE)
• Si aucune expérience pratique prouvée → Score MAX = 30%
• Si expérience <1 an → Score MAX = 50%
• Si niveau junior évident → Score MAX = 60%

Étape 2C - Calcul Final:
Score_Domaine = (Score_Brut × Poids_Domaine) / 100

EXEMPLE DÉTAILLÉ:
Domaine: ".NET Development" - Poids: 40%
CV candidat: AUCUNE mention .NET, seulement Java
→ Score_Brut = 0%
→ Score_Domaine = (0 × 40) / 100 = 0 points
→ Commentaire: "❌ Stack incompatible - profil Java exclusif"

🔴 VÉRIFICATION FINALE OBLIGATOIRE:
Somme_Scores = Σ(tous les Score_Domaine)
Si Somme_Scores ≠ score_matching → ERREUR CRITIQUE → RECALCULE

ÉTAPE 3 - COMMENTAIRE PAR DOMAINE (30-50 mots):
- Utilise ❌ (0-30%), ⚠️ (30-70%), ✅ (70-100%)
- Sois FACTUEL et OBJECTIF dans tes commentaires
- Base-toi UNIQUEMENT sur les FAITS présents dans le CV
- Ne fais PAS d'hypothèses optimistes

EXEMPLE:
JD demande: ".NET, C#, Azure, SQL Server"
Candidat a: "Java, AWS, PostgreSQL"

RÉSULTAT:
{{
  "domaines_analyses": [
    {{
      "domaine": "Stack .NET (C#, ASP.NET Core, Entity Framework)",
      "poids": 40,
      "score": 0,
      "score_max": 40,
      "commentaire": "❌ Aucune expérience .NET/C#. Profil Java exclusivement - incompatibilité majeure sur stack principale.",
      "match": "incompatible"
    }},
    {{
      "domaine": "Cloud Microsoft Azure",
      "poids": 20,
      "score": 8,
      "score_max": 20,
      "commentaire": "⚠️ Expérience AWS uniquement. Compétences cloud transférables mais nécessite formation Azure.",
      "match": "partiel"
    }},
    {{
      "domaine": "SQL Server & T-SQL",
      "poids": 15,
      "score": 10,
      "score_max": 15,
      "commentaire": "✅ Maîtrise PostgreSQL et MySQL - compétences SQL transférables à SQL Server.",
      "match": "bon"
    }}
  ],
  "score_matching": 45,
  "synthese_matching": "PARTIAL MATCH (45/100) - This Java senior profile presents significant challenges for a .NET-focused role, though some transferable competencies exist.

KEY STRENGTHS: The candidate brings 8+ years of enterprise software development experience with proven expertise in cloud platforms (AWS/Azure) and database technologies (PostgreSQL, MySQL). Their experience leading technical teams and architecting scalable solutions demonstrates strong senior-level capabilities. The containerization skills (Docker, Kubernetes) mentioned in their current role are highly relevant.

PARTIAL MATCHES: While the candidate's SQL database experience is strong and transferable to SQL Server, their cloud platform knowledge (AWS/Azure fundamentals) provides a foundation that could accelerate learning of Azure-specific services required for this role. Their experience with agile methodologies and team leadership aligns well with the position's requirements.

CRITICAL GAPS: The most significant concern is the complete absence of .NET stack experience (C#, ASP.NET, Entity Framework), which represents 40% of the role's core requirements (0/40 points). The candidate would require substantial retraining on the entire Microsoft technology stack. Additionally, there's no evidence of Azure-specific service experience (Azure Functions, Service Bus, etc.) beyond basic cloud concepts.

RECOMMENDATION: This profile requires major reconversion and is NOT recommended for immediate placement. Consider only if: (1) the client accepts a 3-6 month ramp-up period, (2) candidate demonstrates strong motivation to transition to .NET, and (3) budget allows for extensive training investment. For urgent needs, seek candidates with existing .NET experience."
}}

Fais :

1. ANALYSE PONDÉRÉE OBLIGATOIRE (voir ci-dessus)
2. Une version reformatée et LÉGÈREMENT reformulée du CV (zéro invention)

2b. PROFIL factuel : paragraphe de 4-5 lignes, profil RÉEL du candidat, sans superlatifs. 3-5 technologies réelles en **gras**.

2c. GRAS ULTRA-SÉLECTIF : 3-5 technologies CRITIQUES réellement maîtrisées.

3. N'ajoute AUCUN mot-clé de la JD que le candidat n'a pas (zéro invention)
4. Garde les intitulés fidèles à la réalité
5. EXPÉRIENCES : style NOMINAL uniforme ("Conception de…", "Réalisation de…"), jamais à la 1re personne ni à l'infinitif ; mêmes faits que le CV, bullets courts (1 ligne), 4-6 max

Réponds en JSON STRICT (sans markdown) avec cette structure:
{{
  "domaines_analyses": [
    {{
      "domaine": "Nom domaine technique/fonctionnel (ex: Stack .NET, Cloud Azure)",
      "poids": 40,
      "score": 15,
      "score_max": 40,
      "commentaire": "Explication 30-50 mots avec ❌/⚠️/✅",
      "match": "incompatible|partiel|bon|excellent"
    }}
  ],
  "score_matching": 45,
  "synthese_matching": "CONCISE PROFESSIONAL SUMMARY (1 paragraph, 80-120 words):
  
  Write a single comprehensive paragraph that includes:
  - Match level (Excellent/Strong/Good/Partial/Weak) with the score (X/100)
  - Candidate's years of experience and seniority level
  - Top 2-3 strongest domains that align perfectly with requirements
  - 1-2 areas that are partial matches or transferable skills
  - 1-2 critical gaps if any
  - Brief recommendation (Recommend/Conditional/Not Recommend)
  
  Keep it analytical and professional. Use concrete examples from the CV. Be honest about both strengths and weaknesses. Make it scannable for busy recruiters. ALWAYS WRITE IN ENGLISH.",
  
  "titre_professionnel_enrichi": "TITRE COURT en {language} (3-5 mots max)",
  
  "profil_enrichi": "Profil NARRATIF 5-6 lignes en {language} avec **3-5 technologies clés** en gras",
  
  "mots_cles_a_mettre_en_gras": ["Technologies RÉELLEMENT présentes dans le CV (recoupant éventuellement la JD) - PAS de technos absentes du CV"],
  
  "competences_enrichies": {{
    "Nom Catégorie 1 (3-6 mots max)": [
      "**Compétence réelle** : description courte (100-150 caractères max) UNIQUEMENT à partir de ce qui figure dans le CV. N'ajoute aucun outil ni résultat absent du CV.",
      "**Autre compétence réelle** : description factuelle et concise, sans invention."
    ],
    "Nom Catégorie 2": [
      "Compétence concise..."
    ]
  }},
  
  RÈGLES ULTRA-CRITIQUES pour les compétences (NON-NÉGOCIABLE):
  - Noms de catégories COURTS (3-6 mots max)
  - 5-6 catégories ADAPTÉES à la JD
  - Chaque catégorie: 3-5 compétences MAXIMUM
  - CHAQUE compétence : 2-3 LIGNES MAXIMUM (100-150 caractères) - NE PAS DÉPASSER
  - N'inclus QUE des compétences réellement présentes dans le CV. AUCUNE compétence inventée, même si la JD la demande.
  - Descriptions factuelles et sobres, sans superlatifs ni résultats chiffrés inventés
  - 1-3 technologies RÉELLES en **gras** par compétence
  
  "experiences_enrichies": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre RÉEL du poste (reformulé seulement si fidèle)",
      "responsabilites": [
        "Configuration **Open edX** incluant structuration et intégration avec **SharePoint** pour gestion contenus",
        "Automatisation processus documentaires via **Power Automate** et **Teams** pour améliorer efficacité"
      ],
      "environment": "**Open edX**, **SharePoint**, **Microsoft 365**, Teams, Power Automate, OneDrive, SQL"
    }}
  ],
  
  "formation_enrichie": [
    {{
      "institution": "Institution name in {language}",
      "diplome": "Degree title in {language}",
      "annee": "2018"
    }}
  ],
  
  "projets_enrichis": [
    "Project description 1 in {language}",
    "Project description 2 in {language}"
  ],
  
  FORMAT OBLIGATOIRE (COPIER format compétences):
  - Responsabilités: Technologies **isolées** dans texte normal (ex: "Configuration **Tech1** incluant **Tech2** pour résultats")
  - Environnement: Liste virgules avec 3-5 technologies **critiques** en gras, autres sans
  - JAMAIS phrases entières en gras
  - Maximum 2-3 mots entre **astérisques**
  
  "score_matching": 45,
  "points_forts": ["ALWAYS in English: key strength 1", "ALWAYS in English: key strength 2"]
}}

🌍 CRITICAL LANGUAGE REQUIREMENT:
- 'domaines_analyses' (domain names AND comments) MUST ALWAYS be in ENGLISH
- 'synthese_matching' MUST ALWAYS be in ENGLISH  
- 'points_forts' MUST ALWAYS be in ENGLISH
- Example domain: "SQL Data Extraction and Manipulation" NOT "Extraction de données SQL"
- Example comment: "❌ No demonstrated experience in SQL data extraction..." NOT "❌ Aucune expérience..."
- Example synthesis: "Java senior profile unsuitable for .NET position..." NOT "Profil Java senior inadapté..."

CRITICAL SCORING RULES:
- 'domaines_analyses' MUST be completed with 5-8 domains totaling EXACTLY 100%
- BE STRICT on scoring - don't give points if candidate lacks the skill
- If stack mismatch (Java vs .NET), give 0 points, not 40-50

🔴🔴🔴 VÉRIFICATION FINALE AVANT RÉPONSE (NON-NÉGOCIABLE) 🔴🔴🔴

AVANT de générer ta réponse JSON, tu DOIS:

1️⃣ CALCULER LA SOMME:
   Somme = domaine1.score + domaine2.score + domaine3.score + ... + domaineN.score
   
2️⃣ VÉRIFIER:
   Si Somme ≠ score_matching → ERREUR → RECALCULE TOUT
   
3️⃣ VÉRIFIER LES POIDS:
   Somme_Poids = domaine1.poids + domaine2.poids + ... + domaineN.poids
   Si Somme_Poids ≠ 100 → ERREUR → RECALCULE TOUT
   
4️⃣ DOUBLE-CHECK:
   Pour chaque domaine: vérifie que score ≤ score_max
   Pour chaque domaine: vérifie que score_max = poids

EXEMPLE DE VÉRIFICATION:
Domaine 1: Stack .NET (40%) → 0/40 points
Domaine 2: Cloud Azure (20%) → 8/20 points  
Domaine 3: SQL Server (15%) → 10/15 points
Domaine 4: DevOps (15%) → 5/15 points
Domaine 5: Agile (10%) → 7/10 points

Vérification poids: 40+20+15+15+10 = 100 ✅
Vérification score: 0+8+10+5+7 = 30 ✅
Donc: score_matching = 30 ✅

Si tu trouves une incohérence → RECALCULE TOUT depuis le début

---

JOB DESCRIPTION:
{jd_text}

---

CV ACTUEL:
{cv_text}

---

IMPORTANT FINAL - RÈGLES JSON STRICTES:
- Génère UNIQUEMENT du JSON valide
- PAS de commentaires (// ou /* */)
- PAS de virgules finales (trailing commas)
- PAS de markdown (```json ou ```)
- TOUS les strings doivent utiliser des guillemets doubles ""
- Vérifie que TOUTES les accolades et crochets sont fermés
- Si tu hésites sur un champ, mets une valeur par défaut plutôt qu'une erreur

Réponds UNIQUEMENT avec du JSON pur, sans rien d'autre avant ou après."""

            print(f">>> Calling Claude API for enrichment with timeout=300s...", flush=True)
            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                timeout=300.0,  # 5 minutes max
                messages=[{"role": "user", "content": prompt}]
            )
            print(f">>> Enrichment API call completed successfully", flush=True)
            
            # 📊 Capturer les métadonnées API
            input_tokens = response.usage.input_tokens if hasattr(response, 'usage') else 0
            output_tokens = response.usage.output_tokens if hasattr(response, 'usage') else 0
            total_tokens = input_tokens + output_tokens
            
        except Exception as e:
            print(f">>> ERROR calling anthropic for enrichment: {repr(e)}", flush=True)
            import traceback
            print(f">>> FULL TRACEBACK:\n{traceback.format_exc()}", flush=True)
            return {}
        
        print(f">>> API Response received, extracting text...", flush=True)
        response_text = response.content[0].text.strip()
        print(f">>> Response length: {len(response_text)} characters", flush=True)
        print(f">>> Response preview (first 500 chars):\n{response_text[:500]}", flush=True)
        
        # Nettoyer JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        print(f">>> Attempting to parse JSON...", flush=True)
        
        # 🔧 NOUVEAU: Tentative de parsing avec retry et correction
        enriched = None
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                if attempt == 0:
                    # Première tentative: parsing direct
                    enriched = json.loads(response_text)
                    print(f">>> JSON parsed successfully on first attempt!", flush=True)
                    break
                else:
                    # Tentatives suivantes: demander à Claude de corriger le JSON
                    print(f">>> Retry {attempt}/{max_retries-1}: Asking Claude to fix JSON...", flush=True)
                    
                    fix_prompt = f"""The following JSON is malformed. Please fix it and return ONLY the corrected JSON without any explanation or markdown:

{response_text}

Return the corrected JSON directly:"""
                    
                    fix_response = client.messages.create(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=8000,
                        timeout=300.0,  # Same as main enrichment call
                        messages=[{"role": "user", "content": fix_prompt}]
                    )
                    
                    fixed_text = fix_response.content[0].text.strip()
                    # Nettoyer le JSON corrigé
                    if fixed_text.startswith('```json'):
                        fixed_text = fixed_text[7:]
                    if fixed_text.startswith('```'):
                        fixed_text = fixed_text[3:]
                    if fixed_text.endswith('```'):
                        fixed_text = fixed_text[:-3]
                    fixed_text = fixed_text.strip()
                    
                    enriched = json.loads(fixed_text)
                    print(f">>> JSON successfully fixed and parsed on attempt {attempt}!", flush=True)
                    break
                    
            except json.JSONDecodeError as e:
                print(f"⚠️ Erreur JSON (attempt {attempt + 1}/{max_retries}): {e}", flush=True)
                if attempt == 0:
                    print(f">>> JSON Error position: {e.pos}", flush=True)
                    print(f">>> Problematic section: {response_text[max(0, e.pos-100):e.pos+100]}", flush=True)
                
                if attempt == max_retries - 1:
                    # Dernier essai échoué: retourner dict vide
                    print(f">>> All parsing attempts failed. Returning empty dict.", flush=True)
                    print(f">>> Full response text:\n{response_text}", flush=True)
                    return {}
                else:
                    # Continuer au prochain retry
                    continue
        
        if enriched is None:
            print(f">>> ERROR: enriched is None after all retries", flush=True)
            return {}
        
        # ✅ FIX: Décoder les entités HTML dans tout le contenu enrichi
        import html
        
        def decode_html_in_dict(obj):
            """Décode récursivement les entités HTML dans un dict/list/str"""
            if isinstance(obj, dict):
                return {k: decode_html_in_dict(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [decode_html_in_dict(item) for item in obj]
            elif isinstance(obj, str):
                return html.unescape(obj)
            else:
                return obj
        
        enriched = decode_html_in_dict(enriched)
        print(f"✅ HTML entities decoded (&#x27; → ')", flush=True)
        
        print(f">>> Keys in enriched: {list(enriched.keys())}", flush=True)
        
        # ⏱️ Calculer le temps de traitement
        processing_time = round(time.time() - start_time, 2)
        
        # 💰 Calculer le coût (prix Claude Sonnet 4.5: $3/MTok input, $15/MTok output)
        cost_input = (input_tokens / 1_000_000) * 3.0
        cost_output = (output_tokens / 1_000_000) * 15.0
        total_cost = round(cost_input + cost_output, 4)
        
        # 📈 Ajouter les métadonnées dans le résultat
        enriched['_metadata'] = {
            'processing_time_seconds': processing_time,
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens,
            'estimated_cost_usd': total_cost
        }
        
        print(f"✅ Enrichissement réussi!")
        
        # ✅ FIX: Si on réutilise le matching, merger les résultats
        if reuse_scoring and matching_analysis:
            print(f"   Mode: Réutilisation du matching du Step 1", flush=True)
            # Récupérer les résultats du Step 1
            enriched['score_matching'] = matching_analysis.get('score_matching', 0)
            enriched['domaines_analyses'] = matching_analysis.get('domaines_analyses', [])
            enriched['synthese_matching'] = matching_analysis.get('synthese_matching', '')
            enriched['points_forts'] = matching_analysis.get('points_forts', [])
            print(f"   Score réutilisé: {enriched['score_matching']}/100")
            print(f"   Domaines réutilisés: {len(enriched['domaines_analyses'])}")
        else:
            print(f"   Mode: Calcul complet du matching", flush=True)
            print(f"   Score matching: {enriched.get('score_matching', 0)}/100")
            print(f"   Domaines analysés: {len(enriched.get('domaines_analyses', []))}")
        print(f"   Mots-clés en gras: {len(enriched.get('mots_cles_a_mettre_en_gras', []))}")
        print(f"   ⏱️ Temps de traitement: {processing_time}s")
        print(f"   📊 Tokens: {total_tokens:,} ({input_tokens:,} in + {output_tokens:,} out)")
        print(f"   💰 Coût estimé: ${total_cost}")
        
        if enriched.get('domaines_analyses'):
            print(f"\n   📊 Détail scoring:")
            for domaine in enriched['domaines_analyses']:
                emoji = domaine.get('match', '')
                if emoji == 'incompatible':
                    emoji = '❌'
                elif emoji == 'partiel':
                    emoji = '⚠️'
                elif emoji in ['bon', 'excellent']:
                    emoji = '✅'
                print(f"      {emoji} {domaine.get('domaine', 'N/A')}: {domaine.get('score', 0)}/{domaine.get('score_max', 0)} ({domaine.get('poids', 0)}%)")
        
        # DEBUG: Afficher une responsabilité pour voir le format
        if enriched.get('experiences_enrichies'):
            first_exp = enriched['experiences_enrichies'][0]
            if first_exp.get('responsabilites'):
                print(f"\n🔍 DEBUG - Première responsabilité :")
                print(f"   {first_exp['responsabilites'][0]}")
            if first_exp.get('environment'):
                print(f"\n🔍 DEBUG - Environnement :")
                print(f"   {first_exp['environment']}")
        
        # 🚨 Vérification critique: le dict ne doit pas être vide
        if not enriched:
            print(f">>> WARNING: enriched dict is EMPTY!", flush=True)
            return {}
        
        # Vérifier les clés essentielles
        required_keys = ['score_matching', 'domaines_analyses', 'profil_enrichi']
        missing_keys = [k for k in required_keys if k not in enriched]
        if missing_keys:
            print(f">>> WARNING: Missing critical keys: {missing_keys}", flush=True)
            print(f">>> Available keys: {list(enriched.keys())}", flush=True)
        
        return enriched

    # ========================================
    # MODULE 4 : MAPPING TMC + RICHTEXT
    # ========================================
    
    def mdbold_to_richtext(self, s: str) -> RichText:
        """Convertit les **bold** markdown en RichText propre sans cascade de gras."""
        import re
        rt = RichText()
        pattern = re.compile(r'\*\*(.*?)\*\*')
        last_end = 0

        # Ajouter le texte avant chaque bloc en gras
        for match in pattern.finditer(s):
            if match.start() > last_end:
                rt.add(s[last_end:match.start()], bold=False, font='Arial')
            # Le texte entre **...** est en gras
            rt.add(match.group(1), bold=True, font='Arial')
            last_end = match.end()

        # Ajouter le texte après le dernier bloc
        if last_end < len(s):
            rt.add(s[last_end:], bold=False, font='Arial')

        return rt

    def map_to_tmc_structure(self, parsed_cv: Dict[str, Any], enriched_cv: Dict[str, Any], template_lang: str = 'FR') -> Dict[str, Any]:
        """Mapper les données enrichies vers la structure TMC"""
        print("🗺️  Mapping vers structure TMC...")

        # 🔧 DÉCODAGE À LA SOURCE : on neutralise toute entité HTML héritée du modèle
        #    (&amp;, &#x27;, &gt;, ...) AVANT de construire RichText et le contexte.
        #    L'échappement XML final (une seule fois) est fait dans generate_tmc_docx.
        import html as _html_src
        def _deep_unescape(o):
            if isinstance(o, str):
                return _html_src.unescape(o)
            if isinstance(o, dict):
                return {k: _deep_unescape(v) for k, v in o.items()}
            if isinstance(o, list):
                return [_deep_unescape(x) for x in o]
            return o
        parsed_cv = _deep_unescape(parsed_cv)
        enriched_cv = _deep_unescape(enriched_cv)
        
        # 1. PROFIL - Convertir en RichText pour supporter le gras (pas d'échappement)
        profil_brut = enriched_cv.get('profil_enrichi', parsed_cv.get('profil_resume', ''))
        profil = self.mdbold_to_richtext(profil_brut) if profil_brut else ''
        
        # 2. COMPÉTENCES - FORMAT CATÉGORISÉ DÉTAILLÉ
        competences_enrichies = enriched_cv.get('competences_enrichies', {})
        
        # Si competences_enrichies est un dict (nouveau format), l'utiliser directement
        if isinstance(competences_enrichies, dict):
            # Supprimer la clé "NOTE" si présente
            skills_categorized = {k: v for k, v in competences_enrichies.items() if k != 'NOTE' and isinstance(v, list)}
        else:
            # Fallback ancien format (liste simple)
            competences = competences_enrichies if isinstance(competences_enrichies, list) else parsed_cv.get('competences', [])
            skills_categorized = {
                'Compétences techniques': competences[:8] if len(competences) >= 8 else competences,
                'Compétences transversales': competences[8:12] if len(competences) > 8 else []
            }
            # Supprimer les catégories vides
            skills_categorized = {k: v for k, v in skills_categorized.items() if v}
        
        # 🔥 Transformation en RichText pour le formatage (pas d'échappement)
        skills_categorized_doc = []
        for cat, skills in skills_categorized.items():
            rt_cat = RichText()
            rt_cat.add(cat, bold=True)
            rt_skills = [self.mdbold_to_richtext(s) for s in skills if s and str(s).strip()]
            skills_categorized_doc.append((rt_cat, rt_skills))
        
        # 3. EXPÉRIENCES - Texte simple pour les responsabilités, RichText pour environnement
        experiences_enrichies = enriched_cv.get('experiences_enrichies', parsed_cv.get('experiences', []))
        work_experience = []
        
        for exp in experiences_enrichies:
            # GARDER les responsabilités en TEXTE SIMPLE (pas RichText) - pas d'échappement
            responsabilites_text = [r for r in exp.get('responsabilites', []) if r and str(r).strip()]
            
            # Convertir l'environnement en RichText pour le gras - pas d'échappement
            environment_brut = exp.get('environment', '')
            environment_rt = self.mdbold_to_richtext(environment_brut) if environment_brut else ''
            
            work_exp = {
                'period': exp.get('periode', ''),
                'company': exp.get('entreprise', ''),
                'position': exp.get('poste', ''),
                'general_responsibilities': responsabilites_text,  # Texte simple
                'environment': environment_rt
            }
            work_experience.append(work_exp)
        
        # 5. CERTIFICATIONS (avec mapping vers format template)
        # 🧹 Nettoyage : on n'affiche JAMAIS "Not specified" / "Date inconnue" / etc.
        #    Un champ manquant devient une chaîne vide (rien ne s'affiche à la place).
        _JUNK_VALUES = {
            'not specified', 'non spécifié', 'non specifie', 'not specify',
            'date inconnue', 'inconnu', 'inconnue', 'unknown',
            'n/a', 'na', 'none', 'null', ''
        }
        def _clean_field(v):
            v = '' if v is None else str(v).strip()
            return '' if v.lower() in _JUNK_VALUES else v

        certifications_raw = parsed_cv.get('certifications', [])
        certifications = []
        for cert in certifications_raw:
            _cn = _clean_field(cert.get('nom', cert.get('name', '')))
            _ci = _clean_field(cert.get('organisme', cert.get('institution', '')))
            _cy = _clean_field(cert.get('annee', cert.get('year', '')))
            _cc = _clean_field(cert.get('pays', cert.get('country', '')))
            # Ligne affichée : Nom │ Organisme │ Année — les champs vides (et leur
            # séparateur) sont simplement omis. Le pays n'est plus affiché.
            _cert_line = ' │ '.join([p for p in [_cn, _ci, _cy] if p])
            certifications.append({
                'name': _cn, 'institution': _ci, 'year': _cy, 'country': _cc,
                'line': _cert_line,
            })
        
        # 6. PROJETS - ❌ Section "Projets pertinents" retirée (demande Aymeric, juin 2026)
        #    Le template masque automatiquement titre + contenu quand la liste est vide.
        projects = []
        
        # ✅ FIX: Ajouter formation enrichie (traduite)
        formation_enrichie = enriched_cv.get('formation_enrichie', parsed_cv.get('formation', []))
        education = []
        for form in formation_enrichie:
            education.append({
                'institution': _clean_field(form.get('institution', '')),
                'degree': _clean_field(form.get('diplome', '')),
                'graduation_year': _clean_field(form.get('annee', '')),
                'country': _clean_field(form.get('pays', '')),
                'level': '',
                'title': _clean_field(form.get('diplome', ''))
            })
        
        # 7. INFORMATIONS PERSONNELLES
        nom_complet = parsed_cv.get('nom_complet', '')
        
        # Séparer prénom et nom
        parts = nom_complet.split() if nom_complet else []
        if len(parts) >= 2:
            first_name = parts[0]
            last_name = ' '.join(parts[1:])
        elif len(parts) == 1:
            first_name = parts[0]
            last_name = ''
        else:
            first_name = 'Prénom'
            last_name = 'Nom'
        
        titre_professionnel = enriched_cv.get('titre_professionnel_enrichi', parsed_cv.get('titre_professionnel', ''))
        lieu_residence = parsed_cv.get('lieu_residence', 'Montréal, Canada')
        langues_list = parsed_cv.get('langues', ['Français', 'Anglais'])
        
        # Traduire les langues selon le template
        if template_lang == 'FR':
            # Si template FR, traduire de l'anglais vers le français
            langue_map = {
                'English': 'Anglais',
                'French': 'Français',
                'Hebrew': 'Hébreu',
                'Russian': 'Russe',
                'Spanish': 'Espagnol',
                'German': 'Allemand',
                'Italian': 'Italien',
                'Portuguese': 'Portugais',
                'Chinese': 'Chinois',
                'Japanese': 'Japonais',
                'Arabic': 'Arabe'
            }
            langues_list = [langue_map.get(lang, lang) for lang in langues_list]
        
        langues = ', '.join(langues_list)
        
        context = {
            # Pour le header (minuscules) - PAS d'échappement
            'first_name': first_name,
            'last_name': last_name,
            'title': titre_professionnel,
            
            # Pour la page 1 (MAJUSCULES) - PAS d'échappement
            'FIRST_NAME': first_name.upper(),
            'LAST_NAME': last_name.upper(),
            'TITLE': titre_professionnel,
            'RESIDENCY': lieu_residence,
            'LANGUAGES': langues,
            
            # AUSSI en minuscules pour compatibilité template
            'residency': lieu_residence,
            'languages': langues,
            
            # Reste du CV
            'summary': profil,
            'skills_categorized': skills_categorized,
            'skills_categorized_doc': skills_categorized_doc,  # 🔥 Version RichText pour le template
            'work_experience': work_experience,
            'education': education,
            'projects': projects,
            'certifications': certifications
        }
        
        print(f"✅ Mapping terminé!")
        print(f"   Nom: [ANONYMIZED]")
        print(f"   Titre: {titre_professionnel}")
        print(f"   Langues: {langues}")
        print(f"   Profil: RichText généré")
        total_competences = sum(len(v) for v in skills_categorized.values() if isinstance(v, list))
        print(f"   Catégories: {len(skills_categorized)}")
        print(f"   Compétences: {total_competences}")
        print(f"   Expériences: {len(work_experience)}")
        
        return context

    # ========================================
    # MODULE 5 : GÉNÉRATION DOCX TMC
    # ========================================
    
    def find_template_file(self, template_name: str = "TMC_NA_template_FR.docx") -> str:
        """Recherche intelligente du template dans plusieurs emplacements possibles"""
        from pathlib import Path
        
        # Liste exhaustive des endroits possibles
        script_dir = Path(__file__).parent
        possible_paths = [
            Path(template_name),  # Current directory
            script_dir / template_name,  # Script directory
            script_dir.parent / "branding" / "templates" / template_name,  # ../../branding/templates/
            script_dir.parent.parent / "branding" / "templates" / template_name,  # ../../../branding/templates/
            Path.home() / template_name,  # Home directory
            Path.home() / "tmc-cv-optimizer" / "branding" / "templates" / template_name,  # Project in home
            Path("/app/branding/templates") / template_name,  # Render deployment path
            Path("/home/ubuntu/tmc-cv-optimizer/branding/templates") / template_name,  # Ubuntu deployment
        ]
        
        # Chercher dans les variables d'environnement aussi
        env_template_path = os.getenv("TMC_TEMPLATE_PATH")
        if env_template_path:
            possible_paths.insert(0, Path(env_template_path))
        
        print(f"   🔍 Recherche du template: {template_name}")
        
        for path in possible_paths:
            try:
                if path.exists() and path.is_file():
                    print(f"   ✅ Template trouvé: {path.resolve()}")
                    return str(path.resolve())
            except (OSError, PermissionError) as e:
                # Ignorer silencieusement les erreurs de permissions
                continue
        
        # Si pas trouvé, afficher tous les chemins essayés
        print(f"   ❌ Template introuvable: {template_name}")
        print(f"   Chemins testés:")
        for path in possible_paths:
            print(f"      - {path}")
        print(f"\n   💡 Astuce: Définir TMC_TEMPLATE_PATH pour spécifier un emplacement personnalisé")
        raise FileNotFoundError(f"Template TMC introuvable: {template_name}")
    
    def generate_tmc_docx(self, context: Dict[str, Any], output_path: str, template_path: str = "TMC_NA_template_FR.docx"):
        """Générer le CV TMC final avec docxtpl"""
        print(f"📝 Génération du CV TMC: {output_path}")
        
        # 🔍 RECHERCHE INTELLIGENTE DU TEMPLATE (nouvelle fonction robuste)
        final_template_path = self.find_template_file(template_path)
        print(f"   📄 Template: {final_template_path}")
        
        # Créer environnement Jinja2 avec filtre pairwise
        jinja_env = jinja2.Environment()
        
        def pairwise(iterable):
            items = list(iterable)
            result = []
            for i in range(0, len(items), 2):
                if i + 1 < len(items):
                    result.append((items[i], items[i + 1]))
                else:
                    result.append((items[i], ''))
            return result
        
        jinja_env.filters['pairwise'] = pairwise
        
        # 🔥 Ajouter la fonction r pour RichText dans le contexte
        context['r'] = lambda x: x
        
        # 🔧 ÉCHAPPEMENT XML — UNE SEULE FOIS (les entités ont été décodées à la source)
        # L\'environnement Jinja passé à docxtpl n\'auto-échappe PAS : les champs texte
        # simples doivent donc être rendus valides en XML (&, <, >). Comme les entités
        # héritées du modèle ont déjà été décodées dans map_to_tmc_structure, on échappe
        # ICI une seule fois -> fini les "&amp;" et "&#x27;" en double. quote=False garde
        # les apostrophes intactes. Les champs RichText (profil, compétences, environnement)
        # gèrent leur propre échappement et ne sont PAS retouchés ici.
        from html import escape as _xml_escape
        def _xs(v):
            return _xml_escape(v, quote=False) if isinstance(v, str) else v

        for key in ['first_name', 'last_name', 'title', 'FIRST_NAME', 'LAST_NAME',
                    'TITLE', 'residency', 'RESIDENCY', 'languages', 'LANGUAGES']:
            if key in context:
                context[key] = _xs(context[key])

        for exp in context.get('work_experience', []):
            for key in ['period', 'company', 'position']:
                if key in exp:
                    exp[key] = _xs(exp[key])
            if isinstance(exp.get('general_responsibilities'), list):
                exp['general_responsibilities'] = [_xs(r) for r in exp['general_responsibilities']]

        for edu in context.get('education', []):
            for key in ['institution', 'degree', 'graduation_year', 'country', 'level', 'title']:
                if key in edu:
                    edu[key] = _xs(edu[key])

        for cert in context.get('certifications', []):
            for key in ['name', 'institution', 'year', 'country', 'line']:
                if key in cert:
                    cert[key] = _xs(cert[key])

        for proj in context.get('projects', []):
            for key in ['name', 'nom', 'description']:
                if key in proj:
                    proj[key] = _xs(proj[key])

        print(f"   ✅ Caractères XML échappés une seule fois (apostrophes & corrigés)")
        
        # Charger le template TMC
        doc = DocxTemplate(final_template_path)
        
        # Rendre le document
        doc.render(context, jinja_env)
        
        # Sauvegarder
        doc.save(output_path)
        print(f"✅ CV TMC généré avec succès!")

    def insert_skills_matrix_page2(self, cv_path, matrix_path, output_path):
        """Insere la skill matrix en PAGE 2 du CV (apres la page de garde, avant les details).
        Compose : couverture + skill matrix + contenu. Insertion verbatim.
        Convertit la matrice en .docx au besoin (via LibreOffice)."""
        from docx import Document
        from docxcompose.composer import Composer
        from pathlib import Path
        import subprocess
        W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        mp = Path(matrix_path)
        if mp.suffix.lower() != '.docx':
            print("   Conversion de la skill matrix en .docx...", flush=True)
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir',
                            str(mp.parent), str(mp)], check=False, timeout=120,
                           stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            converted = mp.parent / (mp.stem + '.docx')
            if not converted.exists():
                raise RuntimeError("Impossible de convertir la skill matrix en .docx")
            matrix_path = str(converted)

        def page_break_index(body):
            for i, el in enumerate(list(body)):
                if el.tag == W + 'p' and el.findall('.//' + W + 'br[@' + W + 'type="page"]'):
                    return i
            return None

        cover = Document(cv_path)
        content = Document(cv_path)
        body_c, body_d = cover.element.body, content.element.body
        idx = page_break_index(body_c)
        if idx is None:
            idx = len(list(body_c))  # pas de page de garde : matrice tout en haut

        # Couverture = elements AVANT le saut de page (le saut est porte par le 1er
        # paragraphe de contenu, donc on s'arrete juste avant).
        for el in list(body_c)[idx:]:
            if el.tag == W + 'sectPr':
                continue
            body_c.remove(el)
        # Forcer un saut de page a la fin de la couverture -> la matrice sera en page 2
        cover.add_page_break()
        # Contenu = a partir du saut de page (garde son pageBreakBefore -> page suivante)
        for el in list(body_d)[:idx]:
            body_d.remove(el)

        matrix = Document(matrix_path)
        bm = matrix.element.body
        for el in list(bm):
            tag = el.tag.split('}')[-1]
            if tag == 'p':
                txt = ''.join(t.text or '' for t in el.findall('.//' + W + 't'))
                if not txt.strip():
                    bm.remove(el)
                else:
                    break
            else:
                break

        comp = Composer(cover)
        comp.append(matrix)
        comp.append(content)
        comp.save(output_path)
        print("Skill matrix inseree en page 2", flush=True)
        return True

    def generate_ms_cv_3parts(self, tmc_context, skills_matrix_path, output_path, 
                              cover_template="TMC_NA_template_EN_Anonymise_CoverPage.docx",
                              content_template="TMC_NA_template_EN_Anonymise_Content.docx"):
        """
        Génère un CV Morgan Stanley en 3 parties:
        1. Cover page (photo + nom + titre + location + langues)
        2. Skills Matrix (uploadée par le recruteur)
        3. Contenu détaillé (profile + skills + experiences + education)
        
        Args:
            tmc_context: Contexte enrichi du candidat
            skills_matrix_path: Path vers le fichier Skills Matrix uploadé
            output_path: Path pour le fichier final
            cover_template: Template pour la cover page
            content_template: Template pour le contenu détaillé
        
        Returns:
            tuple: (success: bool, output_path: str)
        """
        try:
            from pathlib import Path
            from docxcompose.composer import Composer
            from docx import Document
            import shutil
            
            # Dossier temporaire
            temp_dir = Path("/tmp/cv_optimizer_ms")
            temp_dir.mkdir(exist_ok=True)
            
            # ÉTAPE 1: Générer cover page
            print("🎨 Generating cover page...")
            cover_path = temp_dir / "cover.docx"
            
            # ✅ FIX: Passer seulement le nom du template, find_template_file va le chercher
            print(f"   📄 Using cover template: {cover_template}")
            
            self.generate_tmc_docx(
                tmc_context, 
                str(cover_path), 
                template_path=cover_template  # Juste le nom, pas le chemin complet
            )
            print(f"   ✅ Cover page generated: {cover_path.name}")
            
            # ÉTAPE 2: Merger cover + Skills Matrix
            print("🔗 Merging cover with Skills Matrix...")
            cover_with_skills = temp_dir / "cover_and_skills.docx"
            
            # Charger les deux documents
            cover_doc = Document(str(cover_path))
            skills_doc = Document(skills_matrix_path)
            
            # ✅ V1.3.4.2 FIX: Change table width from fixed to auto to prevent horizontal shift
            print("🔧 Fixing Skills Matrix table width...")
            tables_fixed = fix_table_width_to_auto(skills_doc)
            print(f"   ✅ Fixed {tables_fixed} table(s) to auto width")
            
            # V1.3.4 FIX: Ajuster les marges de la Skills Matrix pour correspondre au template
            # Copier les marges du cover vers skills avant merge
            cover_sections = cover_doc.sections
            skills_sections = skills_doc.sections
            
            if cover_sections and skills_sections:
                # Utiliser les marges du template pour la Skills Matrix
                for section in skills_sections:
                    section.top_margin = cover_sections[0].top_margin
                    section.bottom_margin = cover_sections[0].bottom_margin
                    section.left_margin = cover_sections[0].left_margin
                    section.right_margin = cover_sections[0].right_margin
            
            # V1.3.4.1 FIX: Supprimer les espacements au début de la Skills Matrix
            # Ceci assure que le contenu commence exactement en haut de la page
            from docx.shared import Pt
            from docx.oxml import parse_xml
            
            # Supprimer TOUS les paragraphes vides au début du body XML
            # Travailler directement sur body._element pour avoir l'ordre exact
            body = skills_doc.element.body
            elements_to_remove = []
            
            # Parcourir les éléments dans l'ordre et marquer les paragraphes vides au début
            for elem in body:
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                
                if tag == 'p':  # C'est un paragraphe
                    # Vérifier s'il est vide (pas de texte)
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    text_elems = elem.findall('.//w:t', ns)
                    text_content = ''.join([t.text for t in text_elems if t.text])
                    
                    if not text_content.strip():
                        # Paragraphe vide au début → marquer pour suppression
                        elements_to_remove.append(elem)
                    else:
                        # Premier paragraphe avec texte → arrêter
                        break
                elif tag == 'tbl':
                    # On a atteint une table → arrêter
                    break
            
            # Supprimer les éléments marqués
            for elem in elements_to_remove:
                body.remove(elem)
            
            print(f"   🧹 Removed {len(elements_to_remove)} empty paragraphs from Skills Matrix")
            
            # Réinitialiser le spacing du premier élément restant (si paragraphe)
            if skills_doc.paragraphs:
                first_para = skills_doc.paragraphs[0]
                first_para.paragraph_format.space_before = Pt(0)
                first_para.paragraph_format.space_after = Pt(0)
            
            # Ajouter page break après cover
            cover_doc.add_page_break()
            
            # Merger avec docxcompose
            composer = Composer(cover_doc)
            composer.append(skills_doc)
            
            # Sauvegarder
            composer.save(str(cover_with_skills))
            print(f"   ✅ Cover + Skills Matrix merged")
            
            # ÉTAPE 3: Générer contenu détaillé
            print("📝 Generating detailed content...")
            content_path = temp_dir / "content.docx"
            
            # ✅ FIX: Passer seulement le nom du template
            print(f"   📄 Using content template: {content_template}")
            
            self.generate_tmc_docx(
                tmc_context,
                str(content_path),
                template_path=content_template  # Juste le nom, pas le chemin complet
            )
            print(f"   ✅ Content generated: {content_path.name}")
            
            # ÉTAPE 4: Merger tout ensemble
            print("🔗 Merging everything...")
            
            # Charger cover+skills
            final_doc = Document(str(cover_with_skills))
            
            # Ajouter page break avant content
            final_doc.add_page_break()
            
            # Merger avec content
            final_composer = Composer(final_doc)
            content_doc = Document(str(content_path))
            final_composer.append(content_doc)
            
            # Sauvegarder le document final
            final_composer.save(str(output_path))
            print(f"✅ Final CV saved: {output_path}")
            
            # Nettoyer les fichiers temporaires
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return True, str(output_path)
            
        except Exception as e:
            error_msg = f"Error generating MS CV: {str(e)}"
            print(f"❌ {error_msg}")
            import traceback
            traceback.print_exc()
            return False, error_msg
    def apply_bold_post_processing(self, docx_path: str, keywords: list):
        """Post-traiter le document pour mettre en gras les technologies dans les tableaux"""
        print(f"🎨 Application du gras sur les technologies...")
        
        from docx import Document as DocxDocument
        from docx.shared import RGBColor
        import re
        
        doc = DocxDocument(docx_path)
        modifications = 0
        
        print(f"   Recherche des **mot** dans le document...")
        
        def apply_bold_to_runs(paragraph):
            """Trouve **mot** et met en gras UNIQUEMENT ce mot"""
            text = paragraph.text
            if '**' not in text:
                return 0
            
            changes = 0
            # Pattern pour trouver **mot**
            pattern = re.compile(r'\*\*([^*]+)\*\*')
            
            # Reconstituer le paragraphe avec le bon formatage
            matches = list(pattern.finditer(text))
            if not matches:
                return 0
            
            # Supprimer tous les runs existants
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Reconstruire avec le bon formatage
            last_end = 0
            for match in matches:
                # Texte normal avant
                if match.start() > last_end:
                    run = paragraph.add_run(text[last_end:match.start()])
                    run.bold = False
                    run.font.name = 'Arial'
                
                # Texte en gras
                run = paragraph.add_run(match.group(1))
                run.bold = True
                run.font.name = 'Arial'
                changes += 1
                
                last_end = match.end()
            
            # Texte normal après
            if last_end < len(text):
                run = paragraph.add_run(text[last_end:])
                run.bold = False
                run.font.name = 'Arial'
            
            return changes
        
        # Parcourir TOUS les tableaux (où sont les expériences)
        print("   📋 Traitement des tableaux...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        modifications += apply_bold_to_runs(paragraph)
        
        # Parcourir aussi les paragraphes normaux
        print("   📝 Traitement des paragraphes...")
        for paragraph in doc.paragraphs:
            modifications += apply_bold_to_runs(paragraph)
        
        # Sauvegarder
        doc.save(docx_path)
        if modifications > 0:
            print(f"✅ {modifications} mots mis en gras")
        else:
            print(f"⚠️ Aucun **mot** trouvé")
        
        return modifications
        
def main():
    """Point d'entrée CLI"""
    import argparse
    
    parser = argparse.ArgumentParser(description='TMC Universal CV Enricher')
    parser.add_argument('cv_path', help='Chemin du CV (PDF, Word, etc.)')
    parser.add_argument('jd_path', help='Chemin de la Job Description')
    parser.add_argument('--output', '-o', default='cv_enriched_tmc.docx', help='Fichier de sortie')
    
    args = parser.parse_args()
    
    try:
        enricher = TMCUniversalEnricher()
        
        print("\n🚀 TMC UNIVERSAL CV ENRICHER")
        print("=" * 60)
        
        # MODULE 1: Extraction
        print("\n[1/5] Extraction du CV...")
        cv_text = enricher.extract_cv_text(args.cv_path)
        print(f"      ✅ {len(cv_text)} caractères extraits")
        
        # MODULE 2: Parsing
        print("\n[2/5] Parsing intelligent...")
        parsed_cv = enricher.parse_cv_with_claude(cv_text)
        
        # MODULE 3: Enrichissement
        print("\n[3/5] Enrichissement avec IA...")
        jd_text = enricher.read_job_description(args.jd_path)
        enriched_cv = enricher.enrich_cv_with_prompt(parsed_cv, jd_text)
        
        # MODULE 4: Mapping TMC
        print("\n[4/5] Mapping structure TMC...")
        tmc_context = enricher.map_to_tmc_structure(parsed_cv, enriched_cv)
        
        # MODULE 5: Génération
        print("\n[5/5] Génération CV final...")
        enricher.generate_tmc_docx(tmc_context, args.output)
        
        # POST-PROCESSING: Application du gras
        print("\n[POST] Application du gras sur mots-clés...")
        keywords = enriched_cv.get('mots_cles_a_mettre_en_gras', [])
        print(f"   Mots-clés à mettre en gras: {keywords}")
        
        if keywords:
            result = enricher.apply_bold_post_processing(args.output, keywords)
            if result == 0:
                print("   ⚠️ AUCUN mot-clé n'a été mis en gras!")
                print("   Vérifiez que les mots-clés sont bien dans le CV")
        else:
            print("   ⚠️ Aucun mot-clé retourné par l'IA")
        
        # RÉSUMÉ FINAL
        print("\n" + "=" * 60)
        print("🎉 ENRICHISSEMENT TERMINÉ!")
        print("=" * 60)
        print(f"📊 Score matching: {enriched_cv.get('score_matching', 0)}/100")
        
        # Afficher les domaines analysés
        if enriched_cv.get('domaines_analyses'):
            print(f"\n📊 Analyse par domaine:")
            for domaine in enriched_cv['domaines_analyses']:
                match = domaine.get('match', '')
                emoji = '❌' if match == 'incompatible' else '⚠️' if match == 'partiel' else '✅'
                print(f"   {emoji} {domaine.get('domaine', 'N/A')}: {domaine.get('score', 0)}/{domaine.get('score_max', 0)} pts ({domaine.get('poids', 0)}%)")
                print(f"      → {domaine.get('commentaire', 'N/A')}")
        
        if enriched_cv.get('synthese_matching'):
            print(f"\n💬 Synthèse: {enriched_cv['synthese_matching']}")
        
        print(f"\n💪 Points forts:")
        for pf in enriched_cv.get('points_forts', [])[:3]:
            print(f"   • {pf}")
        print(f"\n📄 Fichier généré: {args.output}")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n❌ ERREUR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
